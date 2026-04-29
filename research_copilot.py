"""
Research Writing Copilot
========================
A single-file Streamlit app powered by OpenAI GPT-4o.
Helps you write across domains with a persistent style profile
and a humanization pipeline to keep writing natural and undetectable.

Setup:
    pip install streamlit openai python-docx
    streamlit run research_copilot.py
"""

import json
import os
import re
import time
from pathlib import Path
from datetime import datetime

import streamlit as st
from openai import OpenAI

# ─────────────────────────────────────────────
# === 1. CONFIG & CONSTANTS ===
# ─────────────────────────────────────────────

APP_DIR = Path.home() / ".research_copilot"
PROFILES_FILE = APP_DIR / "style_profiles.json"
HISTORY_FILE = APP_DIR / "session_history.json"
PREFS_FILE = APP_DIR / "preferences.json"

DOMAINS = [
    "Auto-detect",
    "Academic / Scientific",
    "Journalism / Blogging",
    "Technical / Engineering",
    "Business / Corporate",
    "Legal",
    "Medical / Clinical",
    "Humanities / Literary",
    "General / Other",
]

CITATION_STYLES = ["APA", "MLA", "Chicago", "IEEE", "Harvard", "Vancouver"]

MODEL = "gpt-4o"

# ─────────────────────────────────────────────
# === 2. PROFILE MANAGER ===
# ─────────────────────────────────────────────

def ensure_app_dir():
    APP_DIR.mkdir(parents=True, exist_ok=True)

def load_profiles() -> dict:
    ensure_app_dir()
    if PROFILES_FILE.exists():
        with open(PROFILES_FILE, "r") as f:
            return json.load(f)
    return {}

def save_profiles(profiles: dict):
    ensure_app_dir()
    with open(PROFILES_FILE, "w") as f:
        json.dump(profiles, f, indent=2)

def load_history() -> list:
    ensure_app_dir()
    if HISTORY_FILE.exists():
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return []

def save_history(history: list):
    ensure_app_dir()
    # Keep last 50 sessions
    with open(HISTORY_FILE, "w") as f:
        json.dump(history[-50:], f, indent=2)

def load_prefs() -> dict:
    ensure_app_dir()
    if PREFS_FILE.exists():
        with open(PREFS_FILE, "r") as f:
            return json.load(f)
    return {"citation_style": "APA", "theme": "light"}

def save_prefs(prefs: dict):
    ensure_app_dir()
    with open(PREFS_FILE, "w") as f:
        json.dump(prefs, f, indent=2)

# ─────────────────────────────────────────────
# === 3. OPENAI HELPERS ===
# ─────────────────────────────────────────────

def get_client() -> OpenAI:
    api_key = st.session_state.get("api_key", "")
    if not api_key:
        st.error("Please enter your OpenAI API key in the sidebar.")
        st.stop()
    return OpenAI(api_key=api_key)

def call_gpt(system_prompt: str, user_prompt: str, temperature: float = 0.8, max_tokens: int = 2000) -> str:
    client = get_client()
    try:
        response = client.chat.completions.create(
            model=MODEL,
            temperature=temperature,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"OpenAI API error: {e}")
        return ""

def stream_gpt(system_prompt: str, user_prompt: str, temperature: float = 0.8, placeholder=None):
    """Stream GPT response into a Streamlit placeholder."""
    client = get_client()
    full_text = ""
    try:
        stream = client.chat.completions.create(
            model=MODEL,
            temperature=temperature,
            max_tokens=2500,
            stream=True,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content or ""
            full_text += delta
            if placeholder:
                placeholder.markdown(full_text + "▌")
        if placeholder:
            placeholder.markdown(full_text)
    except Exception as e:
        st.error(f"OpenAI API error: {e}")
    return full_text

# ─────────────────────────────────────────────
# === 4. HUMANIZER ENGINE ===
# ─────────────────────────────────────────────

def build_style_fingerprint(samples: list[str], domain: str) -> dict:
    """Analyze writing samples and extract a style fingerprint."""
    combined = "\n\n---\n\n".join(samples)
    system = """You are a linguistic analyst. Analyze the writing samples provided and 
extract a detailed style fingerprint as JSON. Be specific and precise."""
    
    user = f"""Analyze these writing samples from the {domain} domain and return a JSON object with:
{{
  "avg_sentence_length": "short/medium/long",
  "sentence_variety": "description of how sentence lengths vary",
  "vocabulary_level": "simple/moderate/advanced/technical",
  "tone": "formal/semi-formal/conversational/academic/journalistic",
  "voice": "active/passive/mixed",
  "signature_phrases": ["list", "of", "3-5", "characteristic", "phrases or constructions"],
  "transition_style": "how paragraphs and ideas connect",
  "punctuation_habits": "notable punctuation patterns",
  "paragraph_length": "short/medium/long",
  "argumentation_style": "how claims and evidence are structured",
  "domain_specific_notes": "any domain-specific writing conventions observed"
}}

Writing samples:
{combined[:4000]}

Return ONLY the JSON object, no other text."""
    
    result = call_gpt(system, user, temperature=0.3, max_tokens=800)
    try:
        # Strip markdown fences if present
        result = re.sub(r"```json|```", "", result).strip()
        return json.loads(result)
    except:
        return {"raw_analysis": result}

def detect_domain(topic: str) -> str:
    """Auto-detect the domain from a topic description."""
    system = "You are a domain classifier. Return only the domain name, nothing else."
    user = f"""Classify this writing topic into ONE of these domains:
Academic/Scientific, Journalism/Blogging, Technical/Engineering, Business/Corporate, Legal, Medical/Clinical, Humanities/Literary, General

Topic: {topic}

Return only the domain name."""
    return call_gpt(system, user, temperature=0.1, max_tokens=20)

def humanize_text(text: str, style_fingerprint: dict, domain: str, strength: str = "medium") -> str:
    """Apply the 4-technique humanization pipeline."""
    
    strength_instructions = {
        "light": "Make subtle adjustments only. Preserve most of the original structure.",
        "medium": "Meaningfully rewrite while keeping all the core ideas and information.",
        "strong": "Thoroughly rewrite with strong human voice characteristics.",
    }

    fingerprint_str = json.dumps(style_fingerprint, indent=2) if style_fingerprint else "No style profile available."

    system = """You are an expert writing editor specializing in making AI-generated text 
indistinguishable from natural human writing. You use four techniques:
1. PERPLEXITY INJECTION: Use unexpected but precise word choices that break predictable LLM patterns
2. BURSTINESS CONTROL: Deliberately vary sentence lengths — mix very short sentences with longer complex ones
3. VOICE ANCHORING: Weave in reasoning markers, hedges, and perspective indicators natural to humans
4. VOCABULARY SEEDING: Use domain-specific terminology and idiomatic expressions"""

    user = f"""Rewrite the following text to sound authentically human using the writer's style profile.

WRITER'S STYLE FINGERPRINT:
{fingerprint_str}

DOMAIN: {domain}
HUMANIZATION STRENGTH: {strength_instructions[strength]}

RULES:
- Keep ALL facts, citations, and key arguments intact
- Break up overly uniform sentence lengths
- Replace generic AI phrases (e.g., "It is worth noting", "In conclusion", "Furthermore") with natural transitions
- Add subtle imperfections natural to human writing (occasional rhetorical questions, asides, emphasis)
- Match the writer's vocabulary level and tone from the fingerprint
- Do NOT add new claims or information

TEXT TO HUMANIZE:
{text}

Return only the rewritten text, no commentary."""

    return call_gpt(system, user, temperature=0.85, max_tokens=3000)

# ─────────────────────────────────────────────
# === 5. WRITING FUNCTIONS ===
# ─────────────────────────────────────────────

def generate_outline(topic: str, domain: str, paper_type: str, citation_style: str, style_fp: dict) -> str:
    tone_hint = ""
    if style_fp:
        tone_hint = f"Match this writing style: tone={style_fp.get('tone','')}, voice={style_fp.get('voice','')}."
    
    system = f"You are an expert research writer in the {domain} domain."
    user = f"""Create a detailed research outline for the following:

Topic: {topic}
Paper Type: {paper_type}
Citation Style: {citation_style}
{tone_hint}

Generate a structured outline with:
- Clear section headings
- 2-3 bullet points per section describing what to cover
- Suggested word counts per section
- Notes on what evidence/sources to look for

Format it cleanly with markdown headings."""
    
    return call_gpt(system, user, temperature=0.7, max_tokens=1500)

def draft_section(section_title: str, bullet_points: str, topic: str, domain: str, style_fp: dict, context: str = "") -> str:
    fingerprint_str = json.dumps(style_fp, indent=2) if style_fp else "Write in a clear, natural academic style."
    
    system = f"""You are an expert research writer in the {domain} domain. 
Write in the voice described by this style fingerprint:
{fingerprint_str}

CRITICAL: Write as a human expert would — varied sentences, natural flow, no AI clichés."""
    
    user = f"""Write a full draft of this section for a research paper on "{topic}".

Section: {section_title}
Key points to cover:
{bullet_points}

{"Previous context: " + context[:500] if context else ""}

Write 3-5 paragraphs of substantive content. Be specific and authoritative.
Do NOT use these phrases: "It is worth noting", "Furthermore", "In conclusion", "It should be noted", "This paper will".
Return only the section text."""
    
    return call_gpt(system, user, temperature=0.8, max_tokens=1500)

def suggest_citations(claim: str, domain: str, citation_style: str) -> str:
    system = "You are a research librarian expert in academic citations."
    user = f"""For this claim in a {domain} paper, suggest:
1. The type of sources to look for (journals, books, reports, etc.)
2. Key authors or institutions known for this topic
3. How to format a placeholder citation in {citation_style} style
4. Search terms to use in Google Scholar / databases

Claim: {claim}

Be practical and specific."""
    return call_gpt(system, user, temperature=0.5, max_tokens=600)

def refine_text(text: str, instruction: str, domain: str) -> str:
    system = f"You are an expert editor for {domain} writing."
    user = f"""Apply this specific refinement to the text:

Instruction: {instruction}

Text:
{text}

Return only the refined text."""
    return call_gpt(system, user, temperature=0.7, max_tokens=2000)

# ─────────────────────────────────────────────
# === 6. EXPORT FUNCTIONS ===
# ─────────────────────────────────────────────

def export_to_txt(content: dict) -> str:
    lines = []
    lines.append(f"RESEARCH PAPER DRAFT")
    lines.append(f"Topic: {content.get('topic', '')}")
    lines.append(f"Domain: {content.get('domain', '')}")
    lines.append(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("=" * 60)
    lines.append("")
    if content.get("outline"):
        lines.append("OUTLINE")
        lines.append("-" * 40)
        lines.append(content["outline"])
        lines.append("")
    for i, section in enumerate(content.get("sections", []), 1):
        lines.append(f"SECTION {i}: {section.get('title', '')}")
        lines.append("-" * 40)
        lines.append(section.get("content", ""))
        lines.append("")
    return "\n".join(lines)

def export_to_docx(content: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Title
        title = doc.add_heading(content.get("topic", "Research Paper"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph(
            f"Domain: {content.get('domain', '')} | "
            f"Exported: {datetime.now().strftime('%Y-%m-%d')}"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        
        if content.get("outline"):
            doc.add_heading("Outline", level=1)
            # Add outline as plain text (strip markdown)
            clean_outline = re.sub(r"#{1,6}\s*", "", content["outline"])
            doc.add_paragraph(clean_outline)
            doc.add_page_break()
        
        for section in content.get("sections", []):
            doc.add_heading(section.get("title", "Section"), level=1)
            paragraphs = section.get("content", "").split("\n\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_paragraph("")
        
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except ImportError:
        st.warning("python-docx not installed. Run: pip install python-docx")
        return b""

# ─────────────────────────────────────────────
# === 7. SESSION STATE INIT ===
# ─────────────────────────────────────────────

def init_session():
    defaults = {
        "api_key": "",
        "profiles": load_profiles(),
        "prefs": load_prefs(),
        "active_profile": None,
        "topic": "",
        "domain": "Auto-detect",
        "paper_type": "Research Paper",
        "outline": "",
        "sections": [],  # list of {"title": str, "content": str, "humanized": str}
        "draft_tab_section": 0,
        "history": load_history(),
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

# ─────────────────────────────────────────────
# === 8. UI: SIDEBAR ===
# ─────────────────────────────────────────────

def render_sidebar():
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/null/quill-with-ink.png", width=64)
        st.title("Research Copilot")
        st.caption("Powered by GPT-4o")
        st.divider()

        # API Key
        st.subheader("🔑 API Key")
        api_key = st.text_input(
            "OpenAI API Key",
            value=st.session_state.api_key,
            type="password",
            placeholder="sk-...",
            help="Your key is stored only in memory for this session.",
        )
        if api_key != st.session_state.api_key:
            st.session_state.api_key = api_key

        st.divider()

        # Active Style Profile
        st.subheader("🎨 Style Profile")
        profiles = st.session_state.profiles
        profile_names = ["None (use defaults)"] + list(profiles.keys())
        active_idx = 0
        if st.session_state.active_profile in profiles:
            active_idx = profile_names.index(st.session_state.active_profile)
        
        selected = st.selectbox("Active Profile", profile_names, index=active_idx)
        st.session_state.active_profile = None if selected == "None (use defaults)" else selected

        if st.session_state.active_profile:
            profile = profiles[st.session_state.active_profile]
            st.success(f"✅ Profile: **{st.session_state.active_profile}**")
            st.caption(f"Domain: {profile.get('domain', 'N/A')}")
            st.caption(f"Tone: {profile.get('fingerprint', {}).get('tone', 'N/A')}")

        st.divider()

        # Preferences
        st.subheader("⚙️ Preferences")
        citation = st.selectbox(
            "Citation Style",
            CITATION_STYLES,
            index=CITATION_STYLES.index(st.session_state.prefs.get("citation_style", "APA")),
        )
        if citation != st.session_state.prefs.get("citation_style"):
            st.session_state.prefs["citation_style"] = citation
            save_prefs(st.session_state.prefs)

        st.divider()
        st.caption("Files saved to: `~/.research_copilot/`")

# ─────────────────────────────────────────────
# === 9. UI: TAB 1 - SETUP (Style Profile) ===
# ─────────────────────────────────────────────

def render_setup_tab():
    st.header("🧠 Style Profile Manager")
    st.markdown(
        "Paste samples of **your own writing** so the copilot learns your voice. "
        "The more samples, the better the fingerprint."
    )

    profiles = st.session_state.profiles

    col1, col2 = st.columns([2, 1])
    with col1:
        profile_name = st.text_input("Profile Name", placeholder="e.g. My Academic Voice")
        domain = st.selectbox("Domain", [d for d in DOMAINS if d != "Auto-detect"], key="setup_domain")

    with col2:
        st.markdown("**Existing Profiles**")
        if profiles:
            for name in list(profiles.keys()):
                c1, c2 = st.columns([3, 1])
                c1.markdown(f"📄 {name}")
                if c2.button("🗑️", key=f"del_{name}", help="Delete"):
                    del profiles[name]
                    save_profiles(profiles)
                    st.session_state.profiles = profiles
                    st.rerun()
        else:
            st.caption("No profiles yet.")

    st.markdown("**Writing Samples** (paste 2-5 paragraphs of your own writing per sample)")
    
    sample1 = st.text_area("Sample 1", height=150, placeholder="Paste a paragraph from your own writing...")
    sample2 = st.text_area("Sample 2 (optional)", height=120, placeholder="Another sample for better accuracy...")
    sample3 = st.text_area("Sample 3 (optional)", height=120, placeholder="Optional third sample...")

    if st.button("🔬 Analyze & Save Profile", type="primary", use_container_width=True):
        samples = [s for s in [sample1, sample2, sample3] if s.strip()]
        if not profile_name:
            st.error("Please enter a profile name.")
        elif not samples:
            st.error("Please paste at least one writing sample.")
        else:
            with st.spinner("Analyzing your writing style..."):
                fingerprint = build_style_fingerprint(samples, domain)
            
            profiles[profile_name] = {
                "domain": domain,
                "fingerprint": fingerprint,
                "created": datetime.now().isoformat(),
                "sample_count": len(samples),
            }
            save_profiles(profiles)
            st.session_state.profiles = profiles
            st.session_state.active_profile = profile_name
            st.success(f"✅ Profile **{profile_name}** saved and activated!")
            
            with st.expander("View Style Fingerprint"):
                st.json(fingerprint)

    # Show active profile details
    if st.session_state.active_profile and st.session_state.active_profile in profiles:
        st.divider()
        st.subheader(f"📊 Active Profile: {st.session_state.active_profile}")
        profile = profiles[st.session_state.active_profile]
        st.json(profile.get("fingerprint", {}))

# ─────────────────────────────────────────────
# === 10. UI: TAB 2 - OUTLINE ===
# ─────────────────────────────────────────────

def render_outline_tab():
    st.header("📋 Outline Builder")

    col1, col2, col3 = st.columns(3)
    with col1:
        topic = st.text_input(
            "Research Topic",
            value=st.session_state.topic,
            placeholder="e.g. Impact of microplastics on marine ecosystems",
        )
        st.session_state.topic = topic

    with col2:
        domain = st.selectbox("Domain", DOMAINS, key="outline_domain",
                              index=DOMAINS.index(st.session_state.domain) if st.session_state.domain in DOMAINS else 0)
        st.session_state.domain = domain

    with col3:
        paper_type = st.selectbox(
            "Paper Type",
            ["Research Paper", "Review Article", "Essay", "Report", "Blog Post", "Technical Brief", "Case Study"],
            key="paper_type",
        )
        st.session_state.paper_type = paper_type

    if st.button("✨ Generate Outline", type="primary", use_container_width=True):
        if not topic:
            st.error("Please enter a research topic.")
            return
        
        # Auto-detect domain if needed
        effective_domain = domain
        if domain == "Auto-detect":
            with st.spinner("Detecting domain..."):
                effective_domain = detect_domain(topic)
                st.session_state.domain = effective_domain
                st.info(f"Detected domain: **{effective_domain}**")

        # Get active style fingerprint
        style_fp = {}
        if st.session_state.active_profile and st.session_state.active_profile in st.session_state.profiles:
            style_fp = st.session_state.profiles[st.session_state.active_profile].get("fingerprint", {})

        placeholder = st.empty()
        with st.spinner("Generating outline..."):
            outline = stream_gpt(
                system_prompt=f"You are an expert research writer in {effective_domain}.",
                user_prompt=f"Create a detailed research outline for: {topic}\nPaper type: {paper_type}\nCitation: {st.session_state.prefs.get('citation_style','APA')}\nStyle notes: {json.dumps(style_fp)[:300] if style_fp else 'Clear academic style'}",
                temperature=0.7,
                placeholder=placeholder,
            )
        st.session_state.outline = outline

        # Parse outline into sections for Draft tab
        sections = []
        current_section = None
        for line in outline.split("\n"):
            if line.startswith("## ") or line.startswith("# "):
                if current_section:
                    sections.append(current_section)
                current_section = {"title": line.lstrip("#").strip(), "bullets": "", "content": "", "humanized": ""}
            elif current_section and (line.startswith("- ") or line.startswith("* ")):
                current_section["bullets"] += line + "\n"
        if current_section:
            sections.append(current_section)
        
        if sections:
            st.session_state.sections = sections
            st.success(f"✅ Outline generated with {len(sections)} sections. Go to the **Draft** tab to write each section.")

    if st.session_state.outline:
        st.divider()
        st.subheader("Current Outline")
        edited = st.text_area("Edit outline if needed:", value=st.session_state.outline, height=400)
        if edited != st.session_state.outline:
            st.session_state.outline = edited

# ─────────────────────────────────────────────
# === 11. UI: TAB 3 - DRAFT ===
# ─────────────────────────────────────────────

def render_draft_tab():
    st.header("✍️ Section Drafter")

    if not st.session_state.sections:
        st.info("👆 Generate an outline first in the **Outline** tab, or add sections manually below.")
        if st.button("➕ Add Section Manually"):
            st.session_state.sections.append({"title": "New Section", "bullets": "", "content": "", "humanized": ""})
            st.rerun()
        return

    sections = st.session_state.sections
    style_fp = {}
    if st.session_state.active_profile and st.session_state.active_profile in st.session_state.profiles:
        style_fp = st.session_state.profiles[st.session_state.active_profile].get("fingerprint", {})

    # Section selector
    section_titles = [f"{i+1}. {s['title']}" for i, s in enumerate(sections)]
    selected_idx = st.selectbox("Select Section to Work On", range(len(sections)), 
                                 format_func=lambda i: section_titles[i],
                                 key="section_selector")
    
    section = sections[selected_idx]
    st.divider()

    col1, col2 = st.columns(2)
    with col1:
        section["title"] = st.text_input("Section Title", value=section["title"])
    with col2:
        if st.button("🗑️ Remove Section", help="Remove this section"):
            sections.pop(selected_idx)
            st.session_state.sections = sections
            st.rerun()

    section["bullets"] = st.text_area(
        "Key Points to Cover",
        value=section.get("bullets", ""),
        height=100,
        placeholder="- Main argument\n- Supporting evidence\n- Counter-arguments",
    )

    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("🤖 Draft This Section", type="primary", use_container_width=True):
            effective_domain = st.session_state.domain if st.session_state.domain != "Auto-detect" else "Academic"
            # Build context from previous section
            context = sections[selected_idx - 1].get("content", "") if selected_idx > 0 else ""
            
            placeholder = st.empty()
            content = stream_gpt(
                system_prompt=f"Expert research writer in {effective_domain}. Style: {json.dumps(style_fp)[:400] if style_fp else 'clear academic'}",
                user_prompt=f"""Write section "{section['title']}" for a paper on "{st.session_state.topic}".
Key points: {section['bullets']}
{"Context from previous section: " + context[:300] if context else ""}
Write 3-5 substantive paragraphs. Avoid AI clichés.""",
                temperature=0.8,
                placeholder=placeholder,
            )
            section["content"] = content
            sections[selected_idx] = section
            st.session_state.sections = sections

    with col_b:
        if st.button("➕ Add New Section Below", use_container_width=True):
            sections.insert(selected_idx + 1, {"title": "New Section", "bullets": "", "content": "", "humanized": ""})
            st.session_state.sections = sections
            st.rerun()

    # Edit content
    section["content"] = st.text_area(
        "Section Content (edit freely)",
        value=section.get("content", ""),
        height=300,
        placeholder="Draft will appear here, or write directly...",
    )
    sections[selected_idx] = section
    st.session_state.sections = sections

    # Citation helper
    with st.expander("📚 Citation Helper"):
        claim = st.text_input("Enter a specific claim to get citation suggestions:")
        if st.button("Find Sources") and claim:
            effective_domain = st.session_state.domain if st.session_state.domain != "Auto-detect" else "Academic"
            with st.spinner("Finding citation guidance..."):
                suggestions = suggest_citations(claim, effective_domain, st.session_state.prefs.get("citation_style", "APA"))
            st.markdown(suggestions)

    # Progress tracker
    st.divider()
    completed = sum(1 for s in sections if s.get("content", "").strip())
    st.progress(completed / len(sections), text=f"Sections drafted: {completed}/{len(sections)}")

# ─────────────────────────────────────────────
# === 12. UI: TAB 4 - REFINE & HUMANIZE ===
# ─────────────────────────────────────────────

def render_refine_tab():
    st.header("🔍 Refine & Humanize")
    st.markdown(
        "Apply the **4-technique humanization pipeline** to make your writing "
        "sound natural and authentic."
    )

    sections = st.session_state.sections
    if not sections or not any(s.get("content") for s in sections):
        st.info("Draft some sections first in the **Draft** tab.")
        return

    style_fp = {}
    if st.session_state.active_profile and st.session_state.active_profile in st.session_state.profiles:
        style_fp = st.session_state.profiles[st.session_state.active_profile].get("fingerprint", {})

    # Mode selector
    mode = st.radio(
        "What would you like to refine?",
        ["Single Section", "Entire Paper", "Custom Text"],
        horizontal=True,
    )

    if mode == "Single Section":
        section_titles = [f"{i+1}. {s['title']}" for i, s in enumerate(sections)]
        idx = st.selectbox("Select Section", range(len(sections)), format_func=lambda i: section_titles[i])
        text_to_refine = sections[idx].get("content", "")
        
        st.text_area("Current Content", value=text_to_refine, height=200, disabled=True)

        col1, col2 = st.columns(2)
        with col1:
            strength = st.select_slider("Humanization Strength", ["light", "medium", "strong"], value="medium")
        with col2:
            custom_instruction = st.text_input("Additional instruction (optional)", 
                                                placeholder="e.g. Make it more assertive")

        if st.button("✨ Humanize Section", type="primary", use_container_width=True):
            effective_domain = st.session_state.domain if st.session_state.domain != "Auto-detect" else "Academic"
            with st.spinner("Applying humanization pipeline..."):
                result = humanize_text(text_to_refine, style_fp, effective_domain, strength)
                if custom_instruction:
                    result = refine_text(result, custom_instruction, effective_domain)
            sections[idx]["humanized"] = result
            st.session_state.sections = sections
            
        if sections[idx].get("humanized"):
            st.subheader("Humanized Version")
            edited = st.text_area("Edit humanized text:", value=sections[idx]["humanized"], height=300)
            sections[idx]["humanized"] = edited
            st.session_state.sections = sections
            
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("✅ Accept (use as final content)"):
                    sections[idx]["content"] = sections[idx]["humanized"]
                    sections[idx]["humanized"] = ""
                    st.session_state.sections = sections
                    st.success("Accepted!")
            with col_b:
                if st.button("🔄 Re-run Humanizer"):
                    sections[idx]["humanized"] = ""
                    st.session_state.sections = sections
                    st.rerun()

    elif mode == "Entire Paper":
        st.warning("This will process all sections. May take 1-2 minutes.")
        strength = st.select_slider("Humanization Strength", ["light", "medium", "strong"], value="medium")
        
        if st.button("✨ Humanize All Sections", type="primary", use_container_width=True):
            effective_domain = st.session_state.domain if st.session_state.domain != "Auto-detect" else "Academic"
            progress = st.progress(0)
            for i, section in enumerate(sections):
                if section.get("content", "").strip():
                    with st.spinner(f"Humanizing: {section['title']}..."):
                        sections[i]["humanized"] = humanize_text(
                            section["content"], style_fp, effective_domain, strength
                        )
                        sections[i]["content"] = sections[i]["humanized"]
                        sections[i]["humanized"] = ""
                progress.progress((i + 1) / len(sections))
            st.session_state.sections = sections
            st.success("✅ All sections humanized!")

    else:  # Custom Text
        custom_text = st.text_area("Paste any text to humanize:", height=200)
        strength = st.select_slider("Strength", ["light", "medium", "strong"], value="medium", key="custom_strength")
        
        if st.button("✨ Humanize", type="primary") and custom_text:
            effective_domain = st.session_state.domain if st.session_state.domain != "Auto-detect" else "General"
            with st.spinner("Humanizing..."):
                result = humanize_text(custom_text, style_fp, effective_domain, strength)
            st.subheader("Result")
            st.text_area("Humanized text:", value=result, height=250)

    # Quick refinement tools
    st.divider()
    st.subheader("⚡ Quick Refinement Tools")
    
    col1, col2, col3 = st.columns(3)
    quick_text = st.text_area("Text to refine quickly:", height=150, key="quick_refine_text")
    
    with col1:
        if st.button("📝 Fix Grammar & Flow"):
            if quick_text:
                with st.spinner():
                    r = refine_text(quick_text, "Fix grammar, improve sentence flow, and fix awkward phrasing", st.session_state.domain)
                st.text_area("Result:", value=r, height=150, key="qr1")
    with col2:
        if st.button("💡 Make More Concise"):
            if quick_text:
                with st.spinner():
                    r = refine_text(quick_text, "Make significantly more concise without losing key points", st.session_state.domain)
                st.text_area("Result:", value=r, height=150, key="qr2")
    with col3:
        if st.button("🎯 Strengthen Argument"):
            if quick_text:
                with st.spinner():
                    r = refine_text(quick_text, "Strengthen the argumentation and make claims more precise", st.session_state.domain)
                st.text_area("Result:", value=r, height=150, key="qr3")

# ─────────────────────────────────────────────
# === 13. UI: TAB 5 - EXPORT ===
# ─────────────────────────────────────────────

def render_export_tab():
    st.header("📤 Export")

    sections = st.session_state.sections
    has_content = any(s.get("content") for s in sections)

    if not has_content:
        st.info("Draft some content first to export.")
        return

    # Preview
    st.subheader("Paper Preview")
    total_words = 0
    for section in sections:
        content = section.get("content", "")
        if content:
            words = len(content.split())
            total_words += words
            with st.expander(f"📄 {section['title']} ({words} words)"):
                st.markdown(content)

    st.metric("Total Word Count", f"{total_words:,}")

    # Build content dict
    content_dict = {
        "topic": st.session_state.topic,
        "domain": st.session_state.domain,
        "outline": st.session_state.outline,
        "sections": sections,
    }

    st.divider()
    col1, col2, col3 = st.columns(3)

    with col1:
        txt_content = export_to_txt(content_dict)
        st.download_button(
            "📄 Download .txt",
            data=txt_content,
            file_name=f"{st.session_state.topic[:30] or 'paper'}_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with col2:
        docx_bytes = export_to_docx(content_dict)
        if docx_bytes:
            st.download_button(
                "📝 Download .docx",
                data=docx_bytes,
                file_name=f"{st.session_state.topic[:30] or 'paper'}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.info("Install python-docx for Word export:\n`pip install python-docx`")

    with col3:
        if st.button("💾 Save to History", use_container_width=True):
            entry = {
                "date": datetime.now().isoformat(),
                "topic": st.session_state.topic,
                "domain": st.session_state.domain,
                "section_count": len(sections),
                "word_count": total_words,
            }
            history = st.session_state.history
            history.append(entry)
            save_history(history)
            st.session_state.history = history
            st.success("Saved to history!")

    # History
    if st.session_state.history:
        st.divider()
        st.subheader("📚 Session History")
        for entry in reversed(st.session_state.history[-10:]):
            st.markdown(
                f"**{entry.get('topic', 'Untitled')}** — "
                f"{entry.get('domain', '')} — "
                f"{entry.get('word_count', 0):,} words — "
                f"{entry.get('date', '')[:10]}"
            )

# ─────────────────────────────────────────────
# === 14. MAIN APP ===
# ─────────────────────────────────────────────

def main():
    st.set_page_config(
        page_title="Research Copilot",
        page_icon="✍️",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    # Custom CSS
    st.markdown("""
    <style>
        .stTabs [data-baseweb="tab-list"] { gap: 8px; }
        .stTabs [data-baseweb="tab"] { 
            padding: 8px 20px; 
            border-radius: 8px;
            font-weight: 600;
        }
        .stTextArea textarea { font-family: 'Georgia', serif; line-height: 1.7; }
        .stButton button { border-radius: 8px; }
        div[data-testid="metric-container"] { 
            background: #f0f2f6; 
            border-radius: 8px; 
            padding: 12px;
        }
    </style>
    """, unsafe_allow_html=True)

    init_session()
    render_sidebar()

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "🧠 Setup",
        "📋 Outline",
        "✍️ Draft",
        "🔍 Refine",
        "📤 Export",
    ])

    with tab1:
        render_setup_tab()
    with tab2:
        render_outline_tab()
    with tab3:
        render_draft_tab()
    with tab4:
        render_refine_tab()
    with tab5:
        render_export_tab()


if __name__ == "__main__":
    main()
