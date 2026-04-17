import io
import csv
from typing import Any


DISPLAY_COLUMNS = [
    ("author_year", "Author & Year"),
    ("title", "Title"),
    ("research_context", "Research Context"),
    ("methodology", "Methodology"),
    ("independent_variables", "Independent Variables"),
    ("dependent_variable", "Dependent Variable"),
    ("control_variables", "Control Variables"),
    ("findings", "Key Findings"),
    ("theoretical_contributions", "Theoretical Contributions"),
    ("practical_contributions", "Practical Contributions"),
    ("strengths", "Strengths"),
    ("limitations", "Limitations"),
]


def papers_to_csv(papers: list[dict]) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([col[1] for col in DISPLAY_COLUMNS])
    for p in papers:
        row = [p.get(col[0], "") for col in DISPLAY_COLUMNS]
        writer.writerow(row)
    return output.getvalue().encode("utf-8")


def papers_to_excel(papers: list[dict]) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError("openpyxl not installed.")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "EmpiricX Results"

    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    header_fill = PatternFill(start_color="1a1e28", end_color="1a1e28", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_align = Alignment(vertical="top", wrap_text=True)
    thin_border = Border(
        bottom=Side(style="thin", color="E0E0E0"),
        right=Side(style="thin", color="E0E0E0"),
    )

    for col_idx, (key, label) in enumerate(DISPLAY_COLUMNS, 1):
        cell = ws.cell(row=1, column=col_idx, value=label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    ws.row_dimensions[1].height = 36

    for row_idx, paper in enumerate(papers, 2):
        for col_idx, (key, _) in enumerate(DISPLAY_COLUMNS, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=paper.get(key, ""))
            cell.alignment = cell_align
            cell.border = thin_border
            cell.font = Font(name="Calibri", size=9)
        if row_idx % 2 == 0:
            for col_idx in range(1, len(DISPLAY_COLUMNS) + 1):
                ws.cell(row=row_idx, column=col_idx).fill = PatternFill(
                    start_color="F8F9FF", end_color="F8F9FF", fill_type="solid"
                )

    col_widths = [18, 30, 22, 18, 22, 18, 18, 35, 30, 30, 22, 25]
    for col_idx, width in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    ws.freeze_panes = "A2"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def synthesis_to_docx(synthesis: dict, papers: list[dict]) -> bytes:
    """Generate a professional Word document of the cross-paper synthesis report."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import datetime
    except ImportError:
        raise RuntimeError("python-docx not installed.")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def set_heading_style(para, level):
        para.style = doc.styles[f"Heading {level}"]
        for run in para.runs:
            run.font.name = "Arial"
            if level == 1:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1a, 0x1e, 0x28)
            elif level == 2:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x5b, 0x7f, 0xff)

    def add_section_heading(title):
        para = doc.add_heading(title, level=2)
        set_heading_style(para, 2)
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(6)

    def add_bullet(text):
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(str(text))
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        para.paragraph_format.space_after = Pt(3)

    def add_rule(color="5b7fff"):
        rule_para = doc.add_paragraph()
        pPr = rule_para._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pb.append(bottom)
        pPr.append(pb)

    def shade_para(para, hex_fill):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        pPr.append(shd)

    def shade_cell(cell, hex_fill):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    # Title
    title_para = doc.add_paragraph()
    tr = title_para.add_run("EmpiricX")
    tr.font.name = "Arial"
    tr.font.size = Pt(28)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x5b, 0x7f, 0xff)

    sub_para = doc.add_paragraph()
    sr = sub_para.add_run("Cross-Paper Synthesis Report")
    sr.font.name = "Arial"
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x44, 0x44, 0x55)

    date_para = doc.add_paragraph()
    dr = date_para.add_run(
        f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}  \u00b7  {len(papers)} paper(s) analysed"
    )
    dr.font.name = "Arial"
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    doc.add_paragraph()
    add_rule()
    doc.add_paragraph()

    # Overall summary
    summary = synthesis.get("overall_summary", "")
    if summary:
        h = doc.add_heading("Overview", level=1)
        set_heading_style(h, 1)
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(12)
        shade_para(p, "EEF1FF")
        doc.add_paragraph()

    # Main sections
    sections_cfg = [
        ("common_findings",      "Common Findings"),
        ("conflicting_results",  "Conflicting Results"),
        ("methodology_patterns", "Methodology Patterns"),
        ("research_gaps",        "Research Gaps Identified"),
        ("common_weaknesses",    "Common Weaknesses"),
        ("future_directions",    "Future Research Directions"),
    ]

    for key, title in sections_cfg:
        items = synthesis.get(key, [])
        if not items:
            continue
        add_section_heading(title)
        for item in items:
            add_bullet(item)
        doc.add_paragraph()

    # Underexplored variables
    unexplored = synthesis.get("underexplored_variables", [])
    if unexplored:
        add_section_heading("Underexplored Variables")
        p = doc.add_paragraph(", ".join(unexplored))
        run = p.runs[0] if p.runs else p.add_run(", ".join(unexplored))
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x5b, 0x7f, 0xff)
        doc.add_paragraph()

    # Dominant methodology
    dom = synthesis.get("dominant_methodology", "")
    if dom:
        add_section_heading("Dominant Methodology")
        p = doc.add_paragraph(dom)
        if p.runs:
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(11)
        doc.add_paragraph()

    # Analysed papers table
    if papers:
        doc.add_page_break()
        h = doc.add_heading("Analysed Papers", level=1)
        set_heading_style(h, 1)
        h.paragraph_format.space_after = Pt(10)

        col_w = [Cm(3.5), Cm(5.5), Cm(3.5), Cm(3.5)]
        headers = ["Author & Year", "Title", "Methodology", "Key Finding (brief)"]

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, w in enumerate(col_w):
            table.columns[i].width = w

        hdr_row = table.rows[0]
        for i, h_text in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.width = col_w[i]
            run = cell.paragraphs[0].add_run(h_text)
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(cell, "1a1e28")

        for idx, paper in enumerate(papers):
            findings_text = paper.get("findings", "") or ""
            short_finding = findings_text[:120] + ("..." if len(findings_text) > 120 else "")
            vals = [
                paper.get("author_year", "—"),
                paper.get("title", "—"),
                paper.get("methodology", "—"),
                short_finding,
            ]
            fill = "F0F3FF" if idx % 2 == 0 else "FFFFFF"
            row = table.add_row()
            for i, val in enumerate(vals):
                cell = row.cells[i]
                cell.width = col_w[i]
                run = cell.paragraphs[0].add_run(val)
                run.font.name = "Arial"
                run.font.size = Pt(8.5)
                shade_cell(cell, fill)

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fr = fp.add_run("Generated by EmpiricX \u2014 AI-powered Empirical Research Intelligence")
    fr.font.name = "Arial"
    fr.font.size = Pt(8)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xBB)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
