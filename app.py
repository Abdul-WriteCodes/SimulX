"""
Agent43 — AI-Powered Academic Writing System
Single-file Streamlit application
"""

import streamlit as st
import openai
import json
import re
import io
import time
import datetime
import numpy as np
from supabase import create_client, Client
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tiktoken

# ─────────────────────────────────────────────
#  CONFIG & SECRETS
# ─────────────────────────────────────────────

st.set_page_config(
    page_title="Agent43",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

OPENAI_API_KEY   = st.secrets["OPENAI_API_KEY"]
SUPABASE_URL     = st.secrets["SUPABASE_URL"]
SUPABASE_ANON    = st.secrets["SUPABASE_ANON_KEY"]
APP_PASSWORD     = st.secrets["APP_PASSWORD"]

openai_client    = openai.OpenAI(api_key=OPENAI_API_KEY)
supabase: Client = create_client(SUPABASE_URL, SUPABASE_ANON)

GPT_WRITER       = "gpt-4o"
GPT_FAST         = "gpt-4o-mini"
EMBED_MODEL      = "text-embedding-3-small"

# Cost per 1K tokens (USD) — update as OpenAI pricing changes
PRICING = {
    "gpt-4o":          {"in": 0.000005, "out": 0.000015},
    "gpt-4o-mini":     {"in": 0.00000015, "out": 0.0000006},
    "text-embedding-3-small": {"in": 0.00000002, "out": 0.0},
}

# Similarity bands — calibrated for text-embedding-3-small cosine similarity.
# Genuinely synthesized academic text vs its source materials typically scores
# 0.82–0.94 because it shares domain vocabulary and cited concepts.
# "Original" means the model synthesised ideas rather than paraphrased directly.
SIM_LOW    = 0.88   # below → genuinely original synthesis
SIM_HIGH   = 0.94   # above → likely paraphrase / heavy overlap

# ─────────────────────────────────────────────
#  AGENT REGISTRY
# ─────────────────────────────────────────────

AGENTS = {
    # ── International Business ──────────────────────────────────────────
    "Agent Alpha": {
        "class": "International Business",
        "class_code": "IB",
        "signature": "Institutional Economist",
        "description": "Thinks in systems, power structures, and long-run market dynamics. Grounds every argument in trade theory and institutional frameworks. Favours comparative analysis across markets and challenges assumptions that ignore political economy.",
        "theorists": "Dunning, Porter, North, Williamson, Peng",
        "system_prompt": """You are Agent Alpha — a seasoned expert in International Business with over two decades of deep scholarly and applied experience. Your comprehensive command spans trade theory, foreign direct investment, global strategy, organisational behaviour, political economy, emerging markets, cross-cultural management, and international finance.

Your cognitive signature is that of an Institutional Economist. You instinctively situate every argument within systems of power, institutional frameworks, and long-run market dynamics. You ask: what are the rules of the game, who controls them, and how do they shape outcomes? You are drawn to Dunning's OLI paradigm, North's institutional theory, Porter's diamond model, and Williamson's transaction cost economics — not because you are confined to them, but because they are your intellectual home base. You range freely across the entire discipline when the argument demands it.

Your arguments favour comparative analysis across markets and time periods. You challenge assumptions that ignore political economy, path dependency, or the asymmetric distribution of power between MNEs and host nations. You do not oversimplify.

Your prose is authoritative, precise, and structurally rigorous. Your paragraphs build arguments — each sentence advances a claim, not merely restates it. You embed critique naturally; you do not just describe theory, you interrogate it. Your writing sounds like a senior academic who has also advised governments and corporations.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critique, and argumentative sophistication expected of a first-class postgraduate submission."""
    },

    "Agent Beta": {
        "class": "International Business",
        "class_code": "IB",
        "signature": "Emerging Markets Realist",
        "description": "Challenges Western-centric frameworks. Pushes back on assumptions that ignore political economy. Sharp on market volatility, regulatory arbitrage, currency risk, and Global South complexity.",
        "theorists": "Khanna, Palepu, Peng, Rodrik, Lall",
        "system_prompt": """You are Agent Beta — a seasoned expert in International Business with comprehensive mastery across the entire discipline: trade theory, FDI, global strategy, cross-cultural management, international marketing, political risk, and institutional economics.

Your cognitive signature is that of an Emerging Markets Realist. Your lens asks — always — what does this look like from the Global South, from transitional economies, from markets where institutions are weak and volatility is structural? You are deeply sceptical of frameworks built exclusively on Western multinational experience and you say so, with evidence. Khanna and Palepu on institutional voids, Peng on strategy in transition economies, Rodrik on industrial policy — these are your intellectual anchors. But you command the full discipline and draw on mainstream theory precisely to interrogate and complicate it.

Your arguments are sharp, evidence-grounded, and refuse to oversimplify. You stress-test every theoretical claim against real market conditions — currency risk, regulatory arbitrage, political instability, informal institutions. You write with an analytical edge that makes generic arguments look thin by comparison.

Your prose is incisive and confident. You do not hedge unnecessarily. When you challenge a framework, you do it with citations and counter-evidence, not mere scepticism. Your writing sounds like a scholar who has conducted fieldwork in frontier markets and has little patience for boardroom abstractions.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critique, and argumentative sophistication expected of a first-class postgraduate submission."""
    },

    "Agent Gamma": {
        "class": "International Business",
        "class_code": "IB",
        "signature": "Corporate Strategist",
        "description": "Framework-driven, executive-level clarity. Strong bias toward structured, actionable, evidence-anchored arguments. Thinks in PESTLE, OLI, Uppsala, value chain. Never vague.",
        "theorists": "Porter, Ansoff, Johanson, Vahlne, Ghemawat",
        "system_prompt": """You are Agent Gamma — a seasoned expert in International Business with full command of the discipline: corporate strategy, market entry, FDI theory, institutional economics, cross-cultural management, international finance, supply chain, and political risk.

Your cognitive signature is that of a Corporate Strategist. You think in frameworks — not because frameworks are all there is, but because rigorous structure produces clarity, and clarity produces better decisions. PESTLE, OLI, Uppsala model, Porter's Five Forces, Ghemawat's CAGE — these are tools you wield with precision. You also know their limits and you say so when relevant. Your arguments have a strong bias toward actionable insight; you always connect analysis to implication.

You write with executive-level clarity. Your prose is structured, purposeful, and never wastes a sentence. Your arguments are evidence-anchored and your conclusions are never vague. You write like someone who has briefed boards and advised C-suite strategy — authoritative, direct, and analytically airtight. Underneath the clarity is genuine intellectual depth; you do not sacrifice rigour for readability, you achieve both.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critique, and argumentative sophistication expected of a first-class postgraduate submission."""
    },

    # ── International Marketing ──────────────────────────────────────────
    "Agent Delta": {
        "class": "International Marketing",
        "class_code": "IM",
        "signature": "Cultural Intelligence Theorist",
        "description": "Every argument runs through a cultural lens first. Deep command of Hofstede, Trompenaars, Hall. Flags cultural blind spots others miss. Writes with nuance and richness — never reduces markets to demographics.",
        "theorists": "Hofstede, Trompenaars, Hall, de Mooij, Usunier",
        "system_prompt": """You are Agent Delta — a seasoned expert in International Marketing with comprehensive mastery across the full discipline: consumer behaviour, brand management, digital marketing, market research, segmentation, cross-cultural communication, global campaign strategy, pricing, distribution, and marketing theory.

Your cognitive signature is that of a Cultural Intelligence Theorist. Culture is not a variable you add to your analysis — it is the lens through which all marketing phenomena must first be understood. Hofstede's cultural dimensions, Trompenaars' value orientations, Hall's high-context/low-context framework, de Mooij on cross-cultural consumer behaviour — these anchor your worldview. You never reduce markets to demographics or purchasing power. You ask: what do people value, how do they communicate meaning, and how does that shape how they receive a brand, a campaign, a product?

You flag cultural blind spots with authority and evidence. Your arguments are rich with nuance — you can hold complexity without collapsing it into generalisations. When others write "consumers in Market X prefer Y," you write about why, drawing on cultural theory, anthropological insight, and empirical marketing research.

Your prose is sophisticated and layered but never obscure. You write with the assurance of someone who has run international campaigns and watched culturally tone-deaf strategies fail in real markets.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critique, and argumentative sophistication expected of a first-class postgraduate submission."""
    },

    "Agent Epsilon": {
        "class": "International Marketing",
        "class_code": "IM",
        "signature": "Digital & Behavioural Strategist",
        "description": "Lives at the intersection of consumer psychology, data, and digital ecosystems. References platform dynamics, algorithmic influence, and post-pandemic consumer behaviour shifts.",
        "theorists": "Kahneman, Cialdini, Kotler, Chaffey, Ryan",
        "system_prompt": """You are Agent Epsilon — a seasoned expert in International Marketing with full disciplinary command: digital strategy, consumer psychology, brand theory, cultural analysis, market research, segmentation, pricing, global campaign design, and behavioural economics applied to marketing.

Your cognitive signature is that of a Digital and Behavioural Strategist. You operate at the intersection of how the human mind makes decisions and how digital ecosystems shape, amplify, and exploit those decisions at global scale. Kahneman's dual-process theory, Cialdini's influence principles, platform economics, data-driven segmentation, customer journey mapping — these are your native vocabulary. You are acutely aware of how algorithms, social proof, and digital architecture influence consumer behaviour across cultures and markets.

Your arguments are modern, empirically grounded, and alive to the realities of how people actually behave — not how rational-actor models say they should. You reference post-pandemic consumer shifts, the role of influencer ecosystems, the fragmentation of attention, and the challenge of building trust in digital-first markets. But you range freely across the full discipline when the argument demands classical theory.

Your prose is sharp, analytically driven, and energetic without being superficial. You write like someone who has built and optimised international campaigns with real data in hand.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critique, and argumentative sophistication expected of a first-class postgraduate submission."""
    },

    "Agent Zeta": {
        "class": "International Marketing",
        "class_code": "IM",
        "signature": "Brand Equity Architect",
        "description": "Thinks in brand positioning, perception, and identity across international markets. Argues why brands succeed or fail globally with both strategic and emotional depth. Never shallow on brand theory.",
        "theorists": "Keller, Aaker, de Chernatony, Kapferer, Interbrand",
        "system_prompt": """You are Agent Zeta — a seasoned expert in International Marketing with comprehensive mastery across the full field: brand strategy, consumer behaviour, cultural analysis, digital marketing, global campaign design, market entry, segmentation, pricing, and distribution.

Your cognitive signature is that of a Brand Equity Architect. Your world is brand — how it is built, how it travels across cultures, how it accretes or loses value, how it shapes and is shaped by consumer perception. Keller's brand equity model, Aaker's brand identity framework, de Chernatony on brand management, Kapferer on brand identity prism — these are your intellectual infrastructure. But you are not a brand romantic; you are an analyst who understands that brand equity is a business asset with measurable implications.

You argue why brands succeed or fail in international markets with both strategic precision and perceptual depth. You can discuss semiotics and spreadsheets in the same paragraph without losing coherence. You understand that brand is not logo and colour — it is meaning, trust, and the management of expectations across cultural and competitive contexts.

Your prose is confident, elegant, and intellectually serious. You write like someone who has built global brand strategies and understands the gap between theory and execution.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critique, and argumentative sophistication expected of a first-class postgraduate submission."""
    },

    # ── Health & Social Care (NVQ) ───────────────────────────────────────
    "Agent Eta": {
        "class": "Health & Social Care (NVQ)",
        "class_code": "HSC",
        "signature": "Reflective Practitioner",
        "description": "Everything grounded in practice — real care settings, real ethical tensions, real human dignity. Strong on person-centred care, safeguarding, and duty of care. Warm, accountable, professionally anchored.",
        "theorists": "Schön, Gibbs, Johns, Rogers, Maslow",
        "system_prompt": """You are Agent Eta — a seasoned expert in Health and Social Care with comprehensive professional and academic mastery across the full field: reflective practice, person-centred care, safeguarding, mental health, dementia care, end-of-life care, duty of care, ethical frameworks, legislation and policy, communication, equality and diversity, and professional development.

Your cognitive signature is that of a Reflective Practitioner. You believe that practice without reflection is hollow, and reflection without action is indulgent. You are rooted in Schön's reflective practice, Gibbs' reflective cycle, and Johns' model of structured reflection — not as academic exercises but as genuine tools for improving care. Your arguments are always grounded in practice: real care settings, real ethical tensions, the lived reality of service users and care workers.

You write with warmth and professional accountability. Person-centred care is not a buzzword to you — it is a commitment that runs through every argument you make. You are strong on safeguarding, duty of care, dignity, and the professional standards expected of practitioners. You connect theory to practice with specificity and genuine understanding.

Your prose is clear, professionally warm, and ethically serious. You write like a senior practitioner-educator who has worked in care settings and now teaches others — someone who has seen what good and poor care looks like and writes accordingly.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, professional insight, and reflective sophistication expected of an outstanding NVQ submission."""
    },

    "Agent Theta": {
        "class": "Health & Social Care (NVQ)",
        "class_code": "HSC",
        "signature": "Policy & Systems Analyst",
        "description": "Connects individual practice to wider policy and structural frameworks. Deep command of the Care Act, Mental Capacity Act, NHS Long Term Plan, CQC standards. Rigorous, not cold.",
        "theorists": "Bevan, Ham, Marmot, Titmuss, Laming",
        "system_prompt": """You are Agent Theta — a seasoned expert in Health and Social Care with full command of the discipline: health policy, legislation, systems analysis, safeguarding, person-centred care, reflective practice, mental health frameworks, equality and diversity, professional ethics, and social care reform.

Your cognitive signature is that of a Policy and Systems Analyst. You see health and social care as a system shaped by legislation, funding structures, political will, and institutional culture — and you understand how that system affects the individual service user at the point of care. The Care Act 2014, Mental Capacity Act 2005, NHS Long Term Plan, CQC fundamental standards, Marmot's health equity frameworks — these are your primary reference architecture.

You are rigorous but never cold. You connect legislation and policy to their human implications — what does this law mean for a vulnerable adult? What does this CQC standard require of a practitioner? Your arguments operate at both the structural and the individual level, and you move between them fluently.

Your prose is authoritative, precise, and policy-literate. You write like someone who has worked in health commissioning or social care regulation and understands both the intent and the limitations of policy in practice.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, policy literacy, and analytical sophistication expected of an outstanding NVQ submission."""
    },

    "Agent Iota": {
        "class": "Health & Social Care (NVQ)",
        "class_code": "HSC",
        "signature": "Social Justice Advocate",
        "description": "Always asks: who is being left behind? Strong on intersectionality, health inequalities, marginalised communities, and advocacy. Draws on sociological theory as much as care practice. Writes with conviction and moral clarity.",
        "theorists": "Marmot, Wilkinson, Pickett, Oliver, Williams",
        "system_prompt": """You are Agent Iota — a seasoned expert in Health and Social Care with comprehensive mastery across the discipline: social determinants of health, health inequalities, safeguarding, person-centred practice, mental health, disability, legislation, reflective practice, professional ethics, and community care.

Your cognitive signature is that of a Social Justice Advocate. Every assessment you write is animated by a fundamental question: who is being left behind by this system, this policy, this practice — and why? You draw on Marmot's work on health inequalities, Wilkinson and Pickett on the social gradient, disability rights frameworks, intersectionality theory, and the sociology of health and illness. You situate care practice within structures of power, inequality, and exclusion — not to be polemical, but because that context is essential for understanding what good care actually requires.

Your arguments have conviction and moral clarity. You do not mistake neutrality for objectivity — you recognise that silence on inequality is itself a position. But your moral seriousness is always grounded in evidence and theory; your arguments have teeth because they are well-constructed, not merely passionate.

Your prose is intellectually engaged, socially aware, and ethically committed. You write like a practitioner-scholar who has worked with marginalised communities and has both the theoretical framework and the lived perspective to write about it with authority.

CITATION RULES — CRITICAL:
- Cite ONLY from the reference materials provided by the user in their uploaded documents
- Extract author surnames, publication years, and key claims EXACTLY as they appear in the source materials
- Use Harvard in-text citation format: (Author, Year) or Author (Year) argues that...
- At the end of the write-up, produce a full Harvard reference list using ONLY sources drawn from the uploaded materials
- If a source detail is unclear from the materials, cite what is available — do NOT invent or complete missing details
- Never fabricate authors, years, titles, or publishers under any circumstances

Write with the depth, critical awareness, and social justice perspective expected of an outstanding NVQ submission."""
    },
}

CLASS_AGENTS = {
    "International Business":       ["Agent Alpha", "Agent Beta",  "Agent Gamma"],
    "International Marketing":      ["Agent Delta", "Agent Epsilon","Agent Zeta"],
    "Health & Social Care (NVQ)":   ["Agent Eta",   "Agent Theta", "Agent Iota"],
}

# ─────────────────────────────────────────────
#  SUPABASE SQL (shown to user for setup)
# ─────────────────────────────────────────────

SUPABASE_SQL = """
-- Run this in your Supabase SQL Editor

-- Enable pgvector
create extension if not exists vector;

-- Writings table
create table if not exists writings (
    id            bigserial primary key,
    created_at    timestamptz default now(),
    discipline    text,
    agent_name    text,
    context       text,
    word_count    int,
    output_text   text,
    tokens_in     int,
    tokens_out    int,
    cost_usd      numeric(10,6)
);

-- Embeddings table (for similarity checks against past work)
create table if not exists embeddings (
    id          bigserial primary key,
    writing_id  bigint references writings(id) on delete cascade,
    embedding   vector(1536)
);

-- Cost log table
create table if not exists cost_log (
    id          bigserial primary key,
    created_at  timestamptz default now(),
    feature     text,
    model       text,
    tokens_in   int,
    tokens_out  int,
    cost_usd    numeric(10,6)
);

-- ── RAG: Document chunks table ─────────────────────────────────────────────
-- Stores ~600-word chunks of every uploaded source with their embeddings.
-- Enables semantic retrieval across the full document instead of head+tail only.
create table if not exists doc_chunks (
    id           bigserial primary key,
    created_at   timestamptz default now(),
    doc_name     text not null,
    chunk_index  int  not null,
    chunk_text   text not null,
    embedding    vector(1536)
);

-- Index for fast similarity search
create index if not exists doc_chunks_embedding_idx
    on doc_chunks using ivfflat (embedding vector_cosine_ops)
    with (lists = 100);

-- ── RPC: Semantic chunk retrieval ──────────────────────────────────────────
-- Called by the RAG extractor to fetch the most relevant chunks per document.
create or replace function match_doc_chunks(
    query_embedding  vector(1536),
    doc_name_filter  text,
    match_count      int default 6
)
returns table (
    chunk_index  int,
    chunk_text   text,
    similarity   float
)
language sql stable
as $$
    select
        chunk_index,
        chunk_text,
        1 - (embedding <=> query_embedding) as similarity
    from doc_chunks
    where doc_name = doc_name_filter
    order by embedding <=> query_embedding
    limit match_count;
$$;
"""

# ─────────────────────────────────────────────
#  UTILITIES
# ─────────────────────────────────────────────

def calc_cost(model: str, tokens_in: int, tokens_out: int) -> float:
    p = PRICING.get(model, {"in": 0, "out": 0})
    return (tokens_in * p["in"]) + (tokens_out * p["out"])

def count_tokens(text: str, model: str = "gpt-4o") -> int:
    try:
        enc = tiktoken.encoding_for_model(model)
        return len(enc.encode(text))
    except Exception:
        return len(text) // 4

def log_cost(feature: str, model: str, tokens_in: int, tokens_out: int, cost: float):
    try:
        supabase.table("cost_log").insert({
            "feature": feature,
            "model": model,
            "tokens_in": tokens_in,
            "tokens_out": tokens_out,
            "cost_usd": cost
        }).execute()
    except Exception:
        pass

def _extract_raw_text(uploaded_file) -> str:
    """Extract full raw text from a file — no truncation."""
    name = uploaded_file.name.lower()
    text = ""
    try:
        if name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        elif name.endswith(".docx"):
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif name.endswith(".txt"):
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        else:
            text = uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[Could not extract text from {uploaded_file.name}: {e}]"
    return text.strip()


def _chunk_text(text: str, chunk_size: int = 600, overlap: int = 80) -> list:
    """
    Split text into overlapping word-level chunks.
    overlap ensures citations that straddle a boundary are not lost.
    """
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += chunk_size - overlap
    return chunks


def extract_text_from_file(uploaded_file) -> str:
    """
    RAG-powered extraction.

    Strategy:
      1. Extract full text from the document (no truncation).
      2. Split into overlapping ~600-word chunks.
      3. Embed every chunk and store in Supabase `doc_chunks` table.
      4. Run semantic queries for content types most likely to hold
         citable material (arguments, methodology, findings, theory).
      5. Always prepend the first chunk (abstract/intro) and append
         the last chunk (bibliography) so reference metadata is guaranteed.
      6. Return a deduplicated ~15,000-char context window built from
         the highest-relevance chunks across the whole document.

    Falls back gracefully to the old front+back slice if pgvector
    storage or retrieval fails (e.g. table not yet created).
    """
    raw = _extract_raw_text(uploaded_file)
    if not raw or raw.startswith("[Could not"):
        return raw

    chunks = _chunk_text(raw, chunk_size=600, overlap=80)

    # ── Attempt pgvector storage + retrieval ──────────────────────────
    try:
        # Store all chunks with their embeddings
        doc_name = uploaded_file.name
        rows = []
        for idx, chunk in enumerate(chunks):
            emb = get_embedding(chunk)
            rows.append({
                "doc_name":   doc_name,
                "chunk_index": idx,
                "chunk_text": chunk,
                "embedding":  emb
            })
            # Batch insert every 20 rows to avoid request size limits
            if len(rows) == 20:
                supabase.table("doc_chunks").insert(rows).execute()
                rows = []
        if rows:
            supabase.table("doc_chunks").insert(rows).execute()

        # ── Semantic retrieval queries ─────────────────────────────
        # These queries target the content types that carry citable material
        retrieval_queries = [
            "theoretical framework argument literature",
            "methodology research design findings results",
            "critical analysis conclusion implications",
            "author year published argues claims evidence",
        ]

        seen_indices = set()
        ranked_chunks = []

        for query in retrieval_queries:
            q_emb = get_embedding(query)
            # Use pgvector cosine similarity to retrieve top chunks
            result = supabase.rpc(
                "match_doc_chunks",
                {
                    "query_embedding": q_emb,
                    "doc_name_filter": doc_name,
                    "match_count": 6
                }
            ).execute()
            for row in (result.data or []):
                ci = row["chunk_index"]
                if ci not in seen_indices:
                    seen_indices.add(ci)
                    ranked_chunks.append((row.get("similarity", 0), ci, row["chunk_text"]))

        # Always include first chunk (abstract/intro) and last (bibliography)
        for forced_idx, forced_text in [(0, chunks[0]), (len(chunks)-1, chunks[-1])]:
            if forced_idx not in seen_indices:
                seen_indices.add(forced_idx)
                ranked_chunks.append((1.0, forced_idx, forced_text))

        # Sort by chunk index to preserve reading order
        ranked_chunks.sort(key=lambda x: x[1])

        # Build final context — cap at ~15,000 chars
        final_parts = []
        total_chars = 0
        for _, _, text_chunk in ranked_chunks:
            if total_chars + len(text_chunk) > 15000:
                break
            final_parts.append(text_chunk)
            total_chars += len(text_chunk)

        return "\n\n".join(final_parts)

    except Exception:
        # ── Graceful fallback: original front+back slice ───────────
        FRONT = 12000
        BACK  = 6000
        if len(raw) > FRONT + BACK:
            return raw[:FRONT] + "\n\n[...middle section retrieved via fallback...]\n\n" + raw[-BACK:]
        return raw

def get_embedding(text: str) -> list:
    response = openai_client.embeddings.create(
        model=EMBED_MODEL,
        input=text[:8000]
    )
    tokens_used = response.usage.total_tokens
    cost = calc_cost(EMBED_MODEL, tokens_used, 0)
    log_cost("embedding", EMBED_MODEL, tokens_used, 0, cost)
    return response.data[0].embedding

def similarity_score(vec_a: list, vec_b: list) -> float:
    a = np.array(vec_a).reshape(1, -1)
    b = np.array(vec_b).reshape(1, -1)
    return float(cosine_similarity(a, b)[0][0])

def sim_band(score: float) -> tuple:
    if score < SIM_LOW:
        return "🟢 Original", "green"
    elif score < SIM_HIGH:
        return "🟡 Moderate overlap — review recommended", "orange"
    else:
        return "🔴 High similarity — significant overlap detected", "red"

def get_total_cost() -> float:
    try:
        result = supabase.table("cost_log").select("cost_usd").execute()
        return sum(float(r["cost_usd"]) for r in result.data)
    except Exception:
        return 0.0

# ─────────────────────────────────────────────
#  DISPATCHER
# ─────────────────────────────────────────────

def run_dispatcher(context: str) -> dict:
    agent_descriptions = ""
    for name, info in AGENTS.items():
        agent_descriptions += f"\n{name} ({info['class']} — {info['signature']}): {info['description']}\n"

    prompt = f"""You are the Agent43 Dispatcher. Your job is to read an assessment context and recommend the single best agent to write it.

Here are the 9 available agents:
{agent_descriptions}

Assessment context provided by the user:
---
{context}
---

Analyse the context carefully. Consider:
1. Which discipline class does this belong to?
2. Within that class, which agent's cognitive signature best serves this specific brief?

Return ONLY a valid JSON object in exactly this format:
{{
  "recommended": "Agent [Name]",
  "class": "[Class Name]",
  "confidence": "High|Medium|Low",
  "reasoning": "[2-3 sentences explaining why this agent is best suited]",
  "not_alpha": "[one sentence why not — use the actual agent name for that class]",
  "not_beta": "[one sentence why not — use the actual agent name for that class]"
}}

The not_alpha and not_beta keys should explain why the OTHER two agents in the same class are less suited.
Return pure JSON only — no markdown, no explanation outside the JSON."""

    response = openai_client.chat.completions.create(
        model=GPT_FAST,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=600
    )
    raw = response.choices[0].message.content.strip()
    tokens_in  = response.usage.prompt_tokens
    tokens_out = response.usage.completion_tokens
    cost = calc_cost(GPT_FAST, tokens_in, tokens_out)
    log_cost("dispatcher", GPT_FAST, tokens_in, tokens_out, cost)

    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        return json.loads(clean)
    except Exception:
        return {
            "recommended": "Agent Alpha",
            "class": "International Business",
            "confidence": "Low",
            "reasoning": "Could not parse dispatcher response. Defaulting to Agent Alpha.",
            "not_alpha": "N/A",
            "not_beta": "N/A"
        }

# ─────────────────────────────────────────────
#  WRITER
# ─────────────────────────────────────────────

def _parse_structure_sections(structure: str) -> list:
    """Parse user structure into a list of section headings."""
    sections = []
    for line in structure.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        # Strip leading numbers, bullets, dots
        clean = re.sub(r"^[\d]+[\)\.\ ]+|^[-*•]+\s*", "", line).strip()
        if clean:
            sections.append(clean)
    return sections


def _build_source_block(source_texts: list, max_total_chars: int = 80000) -> str:
    if not source_texts:
        return ""
    per_source = max(3000, max_total_chars // len(source_texts))
    block = "\n\n=== REFERENCE MATERIALS ===\n"
    block += "(Cite ONLY from these sources. Do not introduce authors, years, or titles not present here.)\n"
    for i, txt in enumerate(source_texts, 1):
        block += f"\n[Paper {i}]\n{txt[:per_source]}\n"
    block += "\n=== END OF REFERENCE MATERIALS ===\n"
    return block


def _body_word_count(text: str) -> int:
    parts = re.split(
        r"\n(?:References|REFERENCES|Bibliography|BIBLIOGRAPHY)\s*\n",
        text, maxsplit=1
    )
    return len(parts[0].split())


def _strip_references(text: str) -> str:
    parts = re.split(
        r"\n(?:References|REFERENCES|Bibliography|BIBLIOGRAPHY)\s*\n",
        text, maxsplit=1
    )
    return parts[0].rstrip()


def _extract_references(text: str) -> str:
    parts = re.split(
        r"\n((?:References|REFERENCES|Bibliography|BIBLIOGRAPHY)\s*\n)",
        text, maxsplit=1
    )
    if len(parts) >= 3:
        return parts[1] + parts[2]
    return ""


def _do_api_call(system: str, user: str, max_tokens: int,
                 label: str, tokens_in_acc: int, tokens_out_acc: float,
                 cost_acc: float) -> tuple:
    resp = openai_client.chat.completions.create(
        model=GPT_WRITER,
        messages=[{"role": "system", "content": system},
                  {"role": "user",   "content": user}],
        temperature=0.7,
        max_tokens=max_tokens,
        stream=False
    )
    out   = resp.choices[0].message.content
    t_in  = resp.usage.prompt_tokens
    t_out = resp.usage.completion_tokens
    c     = calc_cost(GPT_WRITER, t_in, t_out)
    log_cost(label, GPT_WRITER, t_in, t_out, c)
    return out, tokens_in_acc + t_in, tokens_out_acc + t_out, cost_acc + c


def _do_fast_call(system: str, user: str, max_tokens: int,
                  label: str, tokens_in_acc: int, tokens_out_acc: float,
                  cost_acc: float) -> tuple:
    """Use GPT-4o-mini for cheap extraction/planning stages."""
    resp = openai_client.chat.completions.create(
        model=GPT_FAST,
        messages=[{"role": "system", "content": system},
                  {"role": "user",   "content": user}],
        temperature=0.2,
        max_tokens=max_tokens,
        stream=False
    )
    out   = resp.choices[0].message.content
    t_in  = resp.usage.prompt_tokens
    t_out = resp.usage.completion_tokens
    c     = calc_cost(GPT_FAST, t_in, t_out)
    log_cost(label, GPT_FAST, t_in, t_out, c)
    return out, tokens_in_acc + t_in, tokens_out_acc + t_out, cost_acc + c


# ── Stage 1: Citation Extraction ──────────────────────────────────────────────

def _extract_citations_from_sources(source_texts: list,
                                     tokens_in_acc: int,
                                     tokens_out_acc: float,
                                     cost_acc: float) -> tuple:
    """
    Stage 1: For each uploaded paper, extract a structured citation index.
    Uses GPT-4o-mini (cheap). Returns a clean citation block the writer can
    cite from directly — no raw paper text needed in the write call.
    """
    all_citations = []
    system = "You are a precise academic citation extractor. Extract only what is explicitly stated in the text."

    for i, txt in enumerate(source_texts, 1):
        # Give the extractor the full front+back slice we already captured
        user = f"""Extract the following from this academic source. Return ONLY what you can find explicitly in the text below. Do not infer or complete missing information.

Return as plain text in this exact format:
SOURCE {i}
Authors: [surname, initials of all authors exactly as written]
Year: [publication year]
Title: [full title exactly as written]
Journal/Publisher: [journal name or publisher exactly as written]
Volume/Issue/Pages: [if present]
Key arguments (3-6 bullet points of the main claims/findings, using specific language from the text):
- [claim 1]
- [claim 2]
...
Directly quotable phrases (2-4 short phrases from the text useful for academic writing):
- "[phrase 1]"
- "[phrase 2]"
...

SOURCE TEXT:
{txt[:15000]}"""

        out, tokens_in_acc, tokens_out_acc, cost_acc = _do_fast_call(
            system, user, 800, f"citation_extract_{i}",
            tokens_in_acc, tokens_out_acc, cost_acc
        )
        all_citations.append(out.strip())

    citation_index = "\n\n".join(all_citations)
    return citation_index, tokens_in_acc, tokens_out_acc, cost_acc


# ── Stage 2: Section Plan ──────────────────────────────────────────────────────

def _build_section_plan(context: str, structure: str, scaffold_str: str,
                         citation_index: str, word_count: int,
                         tokens_in_acc: int, tokens_out_acc: float,
                         cost_acc: float) -> tuple:
    """
    Stage 2: Plan which citations to use in each section.
    Returns a section plan JSON so the writer knows exactly what to cite where.
    """
    system = "You are an expert academic essay planner. Return only valid JSON."
    user = f"""You are planning an academic essay. Given the citation index and essay structure below, produce a writing plan.

ESSAY CONTEXT:
{context}

ESSAY STRUCTURE (sections to write in order):
{scaffold_str}

TOTAL BODY WORD COUNT TARGET: {word_count} words

CITATION INDEX (these are the ONLY sources that may be cited):
{citation_index}

Return a JSON array. Each element represents one section:
[
  {{
    "section": "Section heading exactly as given",
    "target_words": <integer — allocate {word_count} words across all sections proportionally>,
    "key_argument": "One sentence stating the core argument this section must make",
    "citations_to_use": ["SOURCE 1", "SOURCE 3"],
    "angle": "Brief note on critical angle — what to challenge, compare, or interrogate in this section"
  }},
  ...
]

Allocate word counts so they sum to exactly {word_count}.
Every section must have at least one citation assigned.
Return pure JSON only — no markdown, no explanation."""

    out, tokens_in_acc, tokens_out_acc, cost_acc = _do_fast_call(
        system, user, 1200, "section_plan",
        tokens_in_acc, tokens_out_acc, cost_acc
    )

    try:
        clean = re.sub(r"```json|```", "", out).strip()
        plan = json.loads(clean)
    except Exception:
        # Fallback: create a simple equal-split plan
        sections = scaffold_str.strip().split("\n")
        per = word_count // max(len(sections), 1)
        plan = [{"section": s.strip(), "target_words": per,
                 "key_argument": "Develop the argument with evidence from sources.",
                 "citations_to_use": ["SOURCE 1"],
                 "angle": "Critically engage with the theory."} for s in sections if s.strip()]

    return plan, tokens_in_acc, tokens_out_acc, cost_acc


# ── Stage 3: Write section by section ─────────────────────────────────────────

def _build_section_prompt(context: str, rubric: str, citation_index: str,
                          sec: dict, word_count: int, section_count: int) -> tuple:
    """
    Build the (system-irrelevant) user prompt and token budget for one section.
    Extracted so streaming and non-streaming paths share identical prompts.
    Returns: (user_prompt_str, out_tokens_int, heading_str)
    """
    heading      = sec.get("section", "Section")
    target_words = int(sec.get("target_words", word_count // max(section_count, 1)))
    key_arg      = sec.get("key_argument", "")
    cite_refs    = sec.get("citations_to_use", ["SOURCE 1"])
    angle        = sec.get("angle", "")
    rubric_block = f"\nMARKING RUBRIC:\n{rubric.strip()}\n" if rubric.strip() else ""

    relevant_citations = []
    for ref in cite_refs:
        pattern = re.compile(rf"({re.escape(ref)}.*?)(?=\nSOURCE \d+|\Z)", re.DOTALL)
        match = pattern.search(citation_index)
        if match:
            relevant_citations.append(match.group(1).strip())
    cite_block = "\n\n".join(relevant_citations) if relevant_citations else citation_index[:3000]

    low  = int(target_words * 0.95)
    high = int(target_words * 1.05)
    out_tokens = max(800, min(4000, int(target_words * 1.6) + 300))

    user = f"""ESSAY CONTEXT:
{context}
{rubric_block}
Write the section titled: {heading}

THIS SECTION'S TARGET: {target_words} words ({low}–{high} words acceptable).
CORE ARGUMENT FOR THIS SECTION: {key_arg}
CRITICAL ANGLE: {angle}

SOURCES FOR THIS SECTION (cite ONLY from these — use Harvard in-text format):
{cite_block}

WRITING RULES:
- Write the section heading on the first line, then the body paragraphs.
- Do NOT write an introduction or conclusion to the essay here — this is one section only.
- Every paragraph advances the argument. No padding, no description without analysis.
- Apply theory critically: explain it, deploy it, challenge its limits.
- Use Harvard in-text citations: (Author, Year) or Author (Year) states...
- Do NOT use markdown symbols (**, ##, *, __).
- Do NOT include a References list — that comes at the end.
- Write exactly {target_words} words of body text for this section."""

    return user, out_tokens, heading


def _write_sections(agent_name: str, context: str, rubric: str,
                    citation_index: str, section_plan: list,
                    word_count: int,
                    tokens_in_acc: int, tokens_out_acc: float,
                    cost_acc: float,
                    stream_container=None) -> tuple:
    """
    Stage 3: Write each section individually using GPT-4o.

    If stream_container is provided (a st.empty() placeholder), each section
    streams token-by-token into the UI as it is written — no blank wait.
    Token counts are captured from the final usage chunk at end of each stream.
    """
    agent  = AGENTS[agent_name]
    system = agent["system_prompt"]
    n      = len(section_plan)

    sections_text  = []
    streamed_so_far = ""   # accumulates full display text for live render

    for i, sec in enumerate(section_plan):
        user, out_tokens, heading = _build_section_prompt(
            context, rubric, citation_index, sec, word_count, n
        )

        if stream_container is not None:
            # ── Streaming path (compatible with openai>=1.0) ────────
            section_buf = ""
            t_in, t_out = 0, 0
            stream_resp = openai_client.chat.completions.create(
                model=GPT_WRITER,
                messages=[{"role": "system", "content": system},
                          {"role": "user",   "content": user}],
                temperature=0.7,
                max_tokens=out_tokens,
                stream=True,
                stream_options={"include_usage": True},
            )
            for chunk in stream_resp:
                delta = chunk.choices[0].delta.content if chunk.choices else None
                if delta:
                    section_buf += delta
                    live_display = streamed_so_far + section_buf
                    paragraphs = [p.strip() for p in live_display.split("\n\n") if p.strip()]
                    rendered = ""
                    for para in paragraphs:
                        if len(para) < 100 and not para.endswith("."):
                            rendered += f'<p class="write-heading">{para}</p>'
                        else:
                            rendered += f'<p class="write-para">{para}</p>'
                    stream_container.markdown(
                        f'<div class="writing-output">{rendered}</div>',
                        unsafe_allow_html=True
                    )
                # Capture usage from the final chunk
                if hasattr(chunk, "usage") and chunk.usage:
                    t_in  = chunk.usage.prompt_tokens
                    t_out = chunk.usage.completion_tokens

            out = section_buf

        else:
            # ── Non-streaming path (fallback) ──────────────────────
            out, tokens_in_acc, tokens_out_acc, cost_acc = _do_api_call(
                system, user, out_tokens, f"write_section_{i+1}",
                tokens_in_acc, tokens_out_acc, cost_acc
            )
            sections_text.append(out.strip())
            continue

        # Accumulate for streaming path
        c = calc_cost(GPT_WRITER, t_in, t_out)
        log_cost(f"write_section_{i+1}", GPT_WRITER, t_in, t_out, c)
        tokens_in_acc  += t_in
        tokens_out_acc += t_out
        cost_acc       += c
        sections_text.append(out.strip())
        streamed_so_far += out.strip() + "\n\n"

    # ── Final: generate consolidated Harvard reference list ───────────────
    ref_user = f"""Based on the citation index below, produce a complete Harvard reference list for all sources.
Format each entry correctly: Author(s) (Year) Title. Journal/Publisher, Volume(Issue), Pages.
List alphabetically by first author surname.
Title the list: References

CITATION INDEX:
{citation_index}

Return ONLY the reference list — nothing else."""

    if stream_container is not None:
        # Stream the reference list too (compatible with openai>=1.0)
        ref_buf = ""
        t_in, t_out = 0, 0
        stream_resp = openai_client.chat.completions.create(
            model=GPT_FAST,
            messages=[{"role": "system", "content": system},
                      {"role": "user",   "content": ref_user}],
            temperature=0.2,
            max_tokens=1000,
            stream=True,
            stream_options={"include_usage": True},
        )
        for chunk in stream_resp:
            delta = chunk.choices[0].delta.content if chunk.choices else None
            if delta:
                ref_buf += delta
                live_display = streamed_so_far + ref_buf
                paragraphs = [p.strip() for p in live_display.split("\n\n") if p.strip()]
                rendered = ""
                for para in paragraphs:
                    if len(para) < 100 and not para.endswith("."):
                        rendered += f'<p class="write-heading">{para}</p>'
                    else:
                        rendered += f'<p class="write-para">{para}</p>'
                stream_container.markdown(
                    f'<div class="writing-output">{rendered}</div>',
                    unsafe_allow_html=True
                )
            if hasattr(chunk, "usage") and chunk.usage:
                t_in  = chunk.usage.prompt_tokens
                t_out = chunk.usage.completion_tokens

        c = calc_cost(GPT_FAST, t_in, t_out)
        log_cost("reference_list", GPT_FAST, t_in, t_out, c)
        tokens_in_acc  += t_in
        tokens_out_acc += t_out
        cost_acc       += c
        ref_out = ref_buf.strip()
    else:
        ref_out, tokens_in_acc, tokens_out_acc, cost_acc = _do_fast_call(
            system, ref_user, 1000, "reference_list",
            tokens_in_acc, tokens_out_acc, cost_acc
        )
        ref_out = ref_out.strip()

    full_output = "\n\n".join(sections_text) + "\n\n" + ref_out
    return full_output, tokens_in_acc, tokens_out_acc, cost_acc


# ── Main entry point ───────────────────────────────────────────────────────────

def run_writer(agent_name: str, context: str, structure: str,
               rubric: str, word_count: int, source_texts: list,
               stream_container=None) -> tuple:
    """
    3-stage pipeline:
      Stage 1 (gpt-4o-mini): Extract structured citation index from each paper.
      Stage 2 (gpt-4o-mini): Plan sections with citation assignments + word targets.
      Stage 3 (gpt-4o):      Write each section individually with exact sources in hand.
                             Streams token-by-token into stream_container if provided.

    This solves the core problem: the model can no longer fabricate citations
    because it works from an extracted index, not raw paper text it might ignore.
    Each write call is small and focused — no word count drift, no continuation hacks.
    """
    total_in, total_out, total_cost = 0, 0, 0.0

    sections = _parse_structure_sections(structure)
    scaffold_str = "\n".join(f"  {s}" for s in sections) if sections else structure.strip()

    # ── Stage 1: Extract citations ─────────────────────────────────────
    citation_index, total_in, total_out, total_cost = _extract_citations_from_sources(
        source_texts, total_in, total_out, total_cost
    )

    # ── Stage 2: Build section plan ────────────────────────────────────
    section_plan, total_in, total_out, total_cost = _build_section_plan(
        context, structure, scaffold_str, citation_index, word_count,
        total_in, total_out, total_cost
    )

    # ── Stage 3: Write section by section (with optional streaming) ─────
    output, total_in, total_out, total_cost = _write_sections(
        agent_name, context, rubric, citation_index, section_plan, word_count,
        total_in, total_out, total_cost,
        stream_container=stream_container
    )

    return output, total_in, total_out, total_cost

# ─────────────────────────────────────────────
#  CITATION VERIFICATION LOOP
# ─────────────────────────────────────────────

def run_citation_verification(output_text: str) -> dict:
    """
    Stage 4 (optional): Cross-check in-text citations against the reference list.

    Finds:
      - Orphaned in-text citations: (Author, Year) in the body with no matching
        reference list entry.
      - Unused references: entries in the reference list never cited in-text.

    Uses GPT-4o-mini (cheap) — returns a structured dict the UI can render.
    """
    body = _strip_references(output_text)
    refs = _extract_references(output_text)

    if not refs.strip():
        return {
            "status": "no_references",
            "orphaned_intext": [],
            "unused_references": [],
            "summary": "No References section found in the output.",
            "parity": False
        }

    prompt = f"""You are an academic citation auditor. Your job is to check that every in-text citation in the BODY has a corresponding entry in the REFERENCE LIST, and that every reference list entry is actually cited in the body.

BODY TEXT (everything before the References section):
---
{body[:6000]}
---

REFERENCE LIST:
---
{refs[:3000]}
---

Analyse carefully. Return ONLY a valid JSON object:
{{
  "orphaned_intext": [
    {{"citation": "(Author, Year)", "note": "brief note on what is missing"}}
  ],
  "unused_references": [
    {{"reference": "Author (Year) Title...", "note": "cited nowhere in body"}}
  ],
  "summary": "1-2 sentence plain-English verdict on the citation parity",
  "parity": true
}}

Rules:
- "parity" should be true ONLY if there are zero orphaned citations AND zero unused references.
- Be precise — a citation like (Smith, 2020) matches a reference entry for Smith (2020) even if the title differs.
- Return pure JSON only — no markdown, no explanation."""

    response = openai_client.chat.completions.create(
        model=GPT_FAST,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=800
    )
    raw       = response.choices[0].message.content.strip()
    t_in      = response.usage.prompt_tokens
    t_out     = response.usage.completion_tokens
    cost      = calc_cost(GPT_FAST, t_in, t_out)
    log_cost("citation_verification", GPT_FAST, t_in, t_out, cost)

    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        result = json.loads(clean)
        result["status"] = "ok"
        return result
    except Exception:
        return {
            "status": "parse_error",
            "orphaned_intext": [],
            "unused_references": [],
            "summary": "Could not parse verification result.",
            "parity": False
        }


# ─────────────────────────────────────────────
#  RISK ASSESSMENT
# ─────────────────────────────────────────────

def run_risk_assessment(text: str) -> dict:
    prompt = f"""You are an academic integrity analyst calibrated to real-world AI detection standards.
Analyse the following academic text excerpt and assess its AI-detectability risk.

You must score FAIRLY and ACCURATELY. Most well-written, discipline-specific academic prose with
real citations, hedged arguments, and varied sentence lengths should score LOW-to-MEDIUM risk.
Reserve HIGH risk for text that has multiple serious AI tells simultaneously.

Score each dimension 0–20 (0 = very human-like, 20 = very AI-typical):
1. Sentence rhythm: Does length and structure vary naturally? (-) Monotone cadence = higher score
2. Transitional language: Overuse of "furthermore", "it is worth noting", "in conclusion"? Each overused phrase +3
3. Specificity: Are arguments anchored to specific evidence, named theorists, real dates/statistics? Generic claims = higher score
4. Structural variety: Does it follow a rigid mechanical pattern per paragraph? Rigid = higher score
5. Hedging calibration: Are qualifiers appropriate to the claims, or do they feel formulaic?

Sum the 5 scores for a total out of 100.

Calibration guidance:
- 0–35: Low risk — reads as competent academic prose, unlikely to trigger detectors
- 36–60: Medium risk — some AI patterns present, human review advisable  
- 61–100: High risk — multiple strong AI signals, significant revision recommended

IMPORTANT: Academic writing by its nature uses formal register and standard structures.
Do NOT penalise for formal tone, Harvard citations, disciplinary vocabulary, or structured arguments.
Only flag genuine AI tells: repetitive sentence openings, hollow transitions, absence of specificity.

Return ONLY a valid JSON object:
{{
  "risk_level": "Low|Medium|High",
  "score": <integer 0-100>,
  "flags": ["specific flag 1 if any", "specific flag 2 if any"],
  "summary": "2-3 sentence honest assessment"
}}

Text to analyse (first 3000 characters):
---
{text[:3000]}
---

Return pure JSON only — no markdown fences."""

    response = openai_client.chat.completions.create(
        model=GPT_FAST,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=400
    )
    raw        = response.choices[0].message.content.strip()
    tokens_in  = response.usage.prompt_tokens
    tokens_out = response.usage.completion_tokens
    cost       = calc_cost(GPT_FAST, tokens_in, tokens_out)
    log_cost("risk_assessment", GPT_FAST, tokens_in, tokens_out, cost)

    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        return json.loads(clean)
    except Exception:
        return {"risk_level": "Unknown", "score": 0, "flags": [], "summary": "Could not parse risk assessment."}

def run_originality_score(output_text: str, source_texts: list) -> dict:
    """
    Compute an originality score (0–100) measuring how much the essay
    goes beyond paraphrase and contributes genuine synthesis.

    Method:
      - Embedding cosine similarity between output and each source chunk.
      - Originality = 100 − (max_similarity × 100), floored at 0.
      - Bonus: +5 for each source used (max +20) — rewards breadth of synthesis.
      - If no sources: score is based on self-assessment by GPT.
    """
    if not source_texts:
        return {"score": 75, "label": "No sources — cannot compute", "color": "grey"}

    output_emb = get_embedding(output_text[:6000])
    max_sim = 0.0
    for src in source_texts:
        src_emb = get_embedding(src[:6000])
        sim = similarity_score(output_emb, src_emb)
        if sim > max_sim:
            max_sim = sim

    # Base originality: inverse of similarity, scaled to 0-100
    base = max(0.0, 1.0 - max_sim) * 100
    # Breadth bonus: using multiple sources = more synthesis
    breadth_bonus = min(len(source_texts) * 5, 20)
    score = min(100, int(base + breadth_bonus))

    if score >= 70:
        label, color = "High Originality", "green"
    elif score >= 45:
        label, color = "Moderate Synthesis", "orange"
    else:
        label, color = "Low Originality", "red"

    return {"score": score, "label": label, "color": color}


# ─────────────────────────────────────────────
#  SIMILARITY CHECKS
# ─────────────────────────────────────────────

def clean_output_text(text: str) -> str:
    """Strip markdown symbols from AI output for clean plain rendering."""
    # Remove bold/italic markers
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    # Remove heading hashes but keep the text
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove underline-style markdown
    text = re.sub(r'__(.*?)__', r'\1', text)
    # Clean up excessive blank lines (max 2)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def count_body_words(output_text: str) -> int:
    """Count words in body only — exclude the References section."""
    # Split off everything from 'References' heading onward
    ref_pattern = re.split(r'\n(?:References|REFERENCES|Bibliography|BIBLIOGRAPHY)\s*\n', output_text, maxsplit=1)
    body = ref_pattern[0] if len(ref_pattern) > 1 else output_text
    return len(body.split())

def check_source_similarity(output_text: str, source_texts: list) -> float:
    if not source_texts:
        return 0.0
    output_emb = get_embedding(output_text[:6000])
    scores = []
    for src in source_texts:
        src_emb = get_embedding(src[:6000])
        scores.append(similarity_score(output_emb, src_emb))
    return max(scores) if scores else 0.0

def check_history_similarity(output_text: str) -> float:
    try:
        result = supabase.table("embeddings").select("embedding").execute()
        if not result.data:
            return 0.0
        output_emb = get_embedding(output_text[:6000])
        scores = []
        for row in result.data:
            stored = row["embedding"]
            if stored:
                scores.append(similarity_score(output_emb, stored))
        return max(scores) if scores else 0.0
    except Exception:
        return 0.0

def save_writing(agent_name: str, discipline: str, context: str,
                 word_count: int, output_text: str,
                 tokens_in: int, tokens_out: int, cost: float):
    try:
        result = supabase.table("writings").insert({
            "discipline":  discipline,
            "agent_name":  agent_name,
            "context":     context[:500],
            "word_count":  word_count,
            "output_text": output_text,
            "tokens_in":   tokens_in,
            "tokens_out":  tokens_out,
            "cost_usd":    cost
        }).execute()
        writing_id = result.data[0]["id"]
        emb = get_embedding(output_text[:6000])
        supabase.table("embeddings").insert({
            "writing_id": writing_id,
            "embedding":  emb
        }).execute()
        return writing_id
    except Exception as e:
        st.warning(f"Could not save to database: {e}")
        return None

# ─────────────────────────────────────────────
#  EXPORT — DOCX
# ─────────────────────────────────────────────

def export_docx(output_text: str, agent_name: str, discipline: str, structure: str) -> bytes:
    doc = Document()

    # Title
    title = doc.add_heading("Agent43 — Academic Write-Up", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Agent: {agent_name}  |  Discipline: {discipline}  |  "
                 f"Generated: {datetime.datetime.now().strftime('%d %B %Y')}").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Split on section headings (numbered or all-caps lines)
    lines = output_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Detect headings: lines starting with number, Roman numeral, or short ALL CAPS
        is_heading = (
            re.match(r"^(\d+[\.\)]|[IVXLC]+\.)\s+", stripped) or
            (stripped.isupper() and len(stripped) < 80) or
            re.match(r"^#{1,3}\s+", stripped)
        )
        if is_heading:
            clean_heading = re.sub(r"^#{1,3}\s+", "", stripped)
            doc.add_heading(clean_heading, level=1)
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ─────────────────────────────────────────────
#  CSS
# ─────────────────────────────────────────────

def inject_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');

    /* ── Base ─────────────────────────────────────────────────────────── */
    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, sans-serif;
    }
    .stApp {
        background: #080c14;
    }
    section[data-testid="stSidebar"] {
        background: #0b0f1a !important;
        border-right: 1px solid rgba(255,255,255,0.05) !important;
    }
    section[data-testid="stSidebar"] > div { padding-top: 0 !important; }

    /* ── Sidebar brand ────────────────────────────────────────────────── */
    .sidebar-brand {
        display: flex;
        align-items: center;
        gap: 0.6rem;
        padding: 1.4rem 1rem 1.2rem;
    }
    .sidebar-logo {
        font-size: 1.4rem;
        color: #6366f1;
        line-height: 1;
    }
    .sidebar-title {
        font-size: 1.05rem;
        font-weight: 700;
        color: #f1f5f9;
        letter-spacing: -0.3px;
    }
    .sidebar-divider {
        height: 1px;
        background: rgba(255,255,255,0.06);
        margin: 0 0.75rem 0.75rem;
    }
    .sidebar-section-label {
        font-size: 0.65rem;
        font-weight: 700;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        color: #475569;
        padding: 0 1rem 0.5rem;
    }
    .sidebar-cost-block {
        padding: 0.75rem 1rem 1rem;
    }
    .sidebar-cost-label {
        font-size: 0.68rem;
        font-weight: 600;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        color: #475569;
        margin-bottom: 0.3rem;
    }
    .sidebar-cost-value {
        font-size: 1.4rem;
        font-weight: 700;
        color: #38bdf8;
        font-family: 'JetBrains Mono', monospace;
        letter-spacing: -0.5px;
    }
    .sidebar-cost-unit {
        font-size: 0.75rem;
        font-weight: 500;
        color: #475569;
        font-family: 'Inter', sans-serif;
    }

    /* ── Sidebar radio nav ────────────────────────────────────────────── */
    div[data-testid="stRadio"] { padding: 0 0.5rem; }
    div[data-testid="stRadio"] label {
        display: block !important;
        padding: 0.55rem 0.75rem !important;
        border-radius: 8px !important;
        font-size: 0.88rem !important;
        font-weight: 500 !important;
        color: #94a3b8 !important;
        cursor: pointer !important;
        transition: background 0.15s, color 0.15s !important;
        margin-bottom: 2px !important;
    }
    div[data-testid="stRadio"] label:hover {
        background: rgba(99,102,241,0.08) !important;
        color: #c7d2fe !important;
    }
    div[data-testid="stRadio"] label[data-checked="true"],
    div[data-testid="stRadio"] input:checked + div {
        background: rgba(99,102,241,0.15) !important;
        color: #a5b4fc !important;
    }
    div[data-testid="stRadio"] > div { gap: 0 !important; }

    /* ── Agent card (sidebar) ─────────────────────────────────────────── */
    .agent-card {
        background: rgba(99,102,241,0.06);
        border: 1px solid rgba(99,102,241,0.15);
        border-radius: 8px;
        padding: 0.7rem 0.9rem;
        margin-bottom: 0.5rem;
    }
    .agent-card-name {
        font-size: 0.82rem;
        font-weight: 600;
        color: #c7d2fe;
        margin-bottom: 0.15rem;
    }
    .agent-card-sig {
        font-size: 0.73rem;
        color: #64748b;
    }

    /* ── Main content area ────────────────────────────────────────────── */
    .main .block-container {
        padding-top: 2rem !important;
        padding-bottom: 4rem !important;
        max-width: 1100px !important;
    }

    /* ── Page header ──────────────────────────────────────────────────── */
    .page-header {
        display: flex;
        align-items: flex-start;
        justify-content: space-between;
        padding: 2rem 2.5rem;
        background: linear-gradient(135deg, #0f172a 0%, #1a1035 50%, #0f1f35 100%);
        border: 1px solid rgba(99,102,241,0.2);
        border-radius: 16px;
        margin-bottom: 2rem;
        position: relative;
        overflow: hidden;
    }
    .page-header::before {
        content: '';
        position: absolute;
        top: -40px; right: -40px;
        width: 200px; height: 200px;
        background: radial-gradient(circle, rgba(99,102,241,0.12) 0%, transparent 70%);
        pointer-events: none;
    }
    .page-header-content h1 {
        font-size: 1.6rem;
        font-weight: 800;
        color: #f1f5f9;
        margin: 0 0 0.3rem;
        letter-spacing: -0.5px;
    }
    .page-header-content p {
        color: #64748b;
        margin: 0;
        font-size: 0.88rem;
    }
    .page-header-badge {
        background: rgba(99,102,241,0.15);
        border: 1px solid rgba(99,102,241,0.3);
        border-radius: 999px;
        padding: 0.35rem 1rem;
        font-size: 0.75rem;
        font-weight: 600;
        color: #a5b4fc;
        white-space: nowrap;
        margin-top: 0.25rem;
    }

    /* ── Section labels ───────────────────────────────────────────────── */
    .section-label {
        font-size: 0.65rem;
        font-weight: 700;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        color: #475569;
        margin-bottom: 0.75rem;
        margin-top: 0.5rem;
    }
    .output-meta {
        font-size: 0.8rem;
        color: #64748b;
        margin-bottom: 0.75rem;
        font-family: 'JetBrains Mono', monospace;
    }

    /* ── Dispatcher card ──────────────────────────────────────────────── */
    .dispatch-card {
        background: linear-gradient(135deg, #052e16 0%, #064e3b 100%);
        border: 1px solid rgba(52,211,153,0.25);
        border-radius: 12px;
        padding: 1.5rem 1.75rem;
        margin: 1rem 0;
        position: relative;
        overflow: hidden;
    }
    .dispatch-card::before {
        content: '';
        position: absolute;
        top: -20px; right: -20px;
        width: 100px; height: 100px;
        background: radial-gradient(circle, rgba(52,211,153,0.1) 0%, transparent 70%);
    }
    .dispatch-card h3 {
        color: #6ee7b7;
        margin: 0 0 0.6rem;
        font-size: 1.15rem;
        font-weight: 700;
    }
    .dispatch-card p {
        color: #a7f3d0;
        margin: 0.3rem 0;
        font-size: 0.87rem;
        line-height: 1.65;
    }

    /* ── Writing output panel ─────────────────────────────────────────── */
    .writing-output {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 2.5rem 3rem;
        color: #1f2937;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    .write-heading {
        font-size: 1rem;
        font-weight: 700;
        color: #111827;
        margin: 1.75rem 0 0.6rem;
        padding-bottom: 0.35rem;
        border-bottom: 2px solid #f3f4f6;
        letter-spacing: -0.2px;
    }
    .write-para {
        font-size: 0.94rem;
        line-height: 1.9;
        color: #374151;
        margin: 0 0 1.1rem;
        text-align: justify;
    }

    /* ── Assessment grid ──────────────────────────────────────────────── */
    .assess-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 1rem;
        margin: 1rem 0;
    }
    .assess-card {
        background: #0d1424;
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 14px;
        padding: 1.4rem 1.6rem;
        transition: border-color 0.2s;
    }
    .assess-card:hover { border-color: rgba(99,102,241,0.2); }
    .assess-card-title {
        font-size: 0.65rem;
        font-weight: 700;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        color: #475569;
        margin-bottom: 0.75rem;
    }
    .assess-score {
        font-size: 2.4rem;
        font-weight: 800;
        letter-spacing: -2px;
        margin-bottom: 0.5rem;
        font-family: 'JetBrains Mono', monospace;
        line-height: 1;
    }
    .assess-label {
        display: inline-flex;
        align-items: center;
        font-size: 0.72rem;
        font-weight: 600;
        padding: 0.25rem 0.85rem;
        border-radius: 999px;
        margin-bottom: 0.85rem;
        letter-spacing: 0.02em;
    }
    .assess-desc {
        font-size: 0.79rem;
        color: #64748b;
        line-height: 1.65;
    }
    .flag-pill {
        display: inline-block;
        background: rgba(245,158,11,0.08);
        color: #f59e0b;
        border: 1px solid rgba(245,158,11,0.2);
        border-radius: 6px;
        padding: 0.18rem 0.6rem;
        font-size: 0.71rem;
        font-weight: 500;
        margin: 0.18rem 0.18rem 0 0;
    }

    /* ── Cost summary bar ─────────────────────────────────────────────── */
    .cost-summary-bar {
        display: flex;
        align-items: center;
        gap: 0.6rem;
        background: #0d1424;
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 10px;
        padding: 0.75rem 1.25rem;
        margin-top: 0.75rem;
        font-size: 0.82rem;
    }
    .cost-summary-label { color: #475569; font-weight: 500; }
    .cost-summary-value {
        color: #38bdf8;
        font-weight: 700;
        font-family: 'JetBrains Mono', monospace;
    }
    .cost-summary-sep { color: #1e293b; }

    /* ── Buttons ──────────────────────────────────────────────────────── */
    div[data-testid="stButton"] > button {
        background: linear-gradient(135deg, #4f46e5 0%, #6d28d9 100%);
        color: #fff;
        border: none;
        border-radius: 10px;
        padding: 0.55rem 1.6rem;
        font-weight: 600;
        font-size: 0.88rem;
        transition: all 0.2s;
        box-shadow: 0 1px 2px rgba(0,0,0,0.2), 0 0 0 1px rgba(99,102,241,0.3);
    }
    div[data-testid="stButton"] > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(99,102,241,0.3), 0 0 0 1px rgba(99,102,241,0.4);
    }
    div[data-testid="stButton"] > button:active { transform: translateY(0); }

    /* Sign out button */
    div[data-testid="stButton"] > button[kind="secondary"] {
        background: transparent;
        border: 1px solid rgba(255,255,255,0.08);
        color: #64748b;
        box-shadow: none;
    }

    /* ── Inputs & text areas ──────────────────────────────────────────── */
    div[data-testid="stTextArea"] textarea,
    div[data-testid="stTextInput"] input {
        background: #0d1424 !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        border-radius: 10px !important;
        color: #e2e8f0 !important;
        font-size: 0.88rem !important;
        transition: border-color 0.2s !important;
    }
    div[data-testid="stTextArea"] textarea:focus,
    div[data-testid="stTextInput"] input:focus {
        border-color: rgba(99,102,241,0.5) !important;
        box-shadow: 0 0 0 3px rgba(99,102,241,0.1) !important;
    }
    div[data-testid="stTextArea"] textarea::placeholder,
    div[data-testid="stTextInput"] input::placeholder {
        color: #334155 !important;
    }
    label[data-testid="stWidgetLabel"] > div > p {
        color: #94a3b8 !important;
        font-size: 0.82rem !important;
        font-weight: 500 !important;
    }

    /* ── Selectbox ────────────────────────────────────────────────────── */
    div[data-testid="stSelectbox"] > div > div {
        background: #0d1424 !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        border-radius: 10px !important;
        color: #e2e8f0 !important;
    }

    /* ── Number input ─────────────────────────────────────────────────── */
    div[data-testid="stNumberInput"] input {
        background: #0d1424 !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        border-radius: 10px !important;
        color: #e2e8f0 !important;
    }

    /* ── File uploader ────────────────────────────────────────────────── */
    div[data-testid="stFileUploader"] {
        background: #0d1424 !important;
        border: 1px dashed rgba(99,102,241,0.25) !important;
        border-radius: 12px !important;
    }
    div[data-testid="stFileUploader"] label { color: #64748b !important; }

    /* ── Expander ─────────────────────────────────────────────────────── */
    div[data-testid="stExpander"] {
        background: #0d1424 !important;
        border: 1px solid rgba(255,255,255,0.06) !important;
        border-radius: 12px !important;
        overflow: hidden;
    }
    div[data-testid="stExpander"] summary {
        color: #94a3b8 !important;
        font-size: 0.85rem !important;
        font-weight: 500 !important;
        padding: 0.75rem 1rem !important;
    }

    /* ── Metrics ──────────────────────────────────────────────────────── */
    div[data-testid="stMetric"] {
        background: #0d1424;
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 12px;
        padding: 1rem 1.25rem;
    }
    div[data-testid="stMetricLabel"] p { color: #64748b !important; font-size: 0.75rem !important; }
    div[data-testid="stMetricValue"]  { color: #f1f5f9 !important; font-weight: 700 !important; }

    /* ── Horizontal rule ──────────────────────────────────────────────── */
    hr { border-color: rgba(255,255,255,0.06) !important; }

    /* ── Info / warning / error / success boxes ───────────────────────── */
    div[data-testid="stAlert"] {
        border-radius: 10px !important;
        border: none !important;
    }

    /* ── Step headings ────────────────────────────────────────────────── */
    h3 { color: #e2e8f0 !important; font-weight: 700 !important; letter-spacing: -0.3px !important; }
    h2 { color: #f1f5f9 !important; font-weight: 800 !important; letter-spacing: -0.5px !important; }

    /* ── Caption ──────────────────────────────────────────────────────── */
    div[data-testid="stCaptionContainer"] p { color: #475569 !important; font-size: 0.8rem !important; }

    /* ── Download buttons ─────────────────────────────────────────────── */
    div[data-testid="stDownloadButton"] > button {
        background: #0d1424 !important;
        color: #94a3b8 !important;
        border: 1px solid rgba(255,255,255,0.1) !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        box-shadow: none !important;
    }
    div[data-testid="stDownloadButton"] > button:hover {
        background: #141d30 !important;
        color: #e2e8f0 !important;
        border-color: rgba(99,102,241,0.3) !important;
        transform: translateY(-1px) !important;
    }

    /* ── Auth screen ──────────────────────────────────────────────────── */
    .auth-screen {
        min-height: 100vh;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
    }
    .auth-card {
        background: #0d1424;
        border: 1px solid rgba(99,102,241,0.2);
        border-radius: 20px;
        padding: 2.5rem;
        width: 360px;
        text-align: center;
        box-shadow: 0 20px 60px rgba(0,0,0,0.4), 0 0 0 1px rgba(99,102,241,0.1);
    }
    .auth-logo {
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
    }
    .auth-title {
        font-size: 1.4rem;
        font-weight: 800;
        color: #f1f5f9;
        margin-bottom: 0.25rem;
        letter-spacing: -0.5px;
    }
    .auth-subtitle {
        font-size: 0.82rem;
        color: #475569;
        margin-bottom: 2rem;
    }

    /* ── Hero ─────────────────────────────────────────────────────────── */
    .hero {
        position: relative;
        overflow: hidden;
        background: linear-gradient(160deg, #080c14 0%, #0e1528 40%, #0a1220 100%);
        border: 1px solid rgba(99,102,241,0.18);
        border-radius: 20px;
        padding: 4rem 3rem 3.5rem;
        margin-bottom: 1.5rem;
        text-align: center;
    }
    .hero-glow {
        position: absolute;
        width: 360px; height: 360px;
        border-radius: 50%;
        filter: blur(80px);
        pointer-events: none;
        opacity: 0.4;
    }
    .hero-glow-left  { background: radial-gradient(circle, #4f46e5, transparent 70%); top: -80px; left: -80px; }
    .hero-glow-right { background: radial-gradient(circle, #0ea5e9, transparent 70%); bottom: -80px; right: -80px; }
    .hero-inner { position: relative; z-index: 1; }
    .hero-eyebrow {
        font-size: 0.65rem;
        font-weight: 700;
        letter-spacing: 0.18em;
        color: #6366f1;
        margin-bottom: 1.1rem;
    }
    .hero-title {
        font-size: 2.6rem !important;
        font-weight: 800 !important;
        color: #f1f5f9 !important;
        letter-spacing: -1.5px !important;
        line-height: 1.15 !important;
        margin: 0 0 1.1rem !important;
    }
    .hero-subtitle {
        font-size: 1rem;
        color: #64748b;
        max-width: 540px;
        margin: 0 auto;
        line-height: 1.7;
    }

    /* ── Capability grid ──────────────────────────────────────────────── */
    .capability-grid {
        display: grid;
        grid-template-columns: 1fr 1fr 1fr;
        gap: 1rem;
        margin-bottom: 1.5rem;
    }
    .cap-card {
        background: #0d1424;
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 14px;
        padding: 1.4rem 1.5rem 1.25rem;
        transition: border-color 0.2s, transform 0.2s;
    }
    .cap-card:hover {
        border-color: rgba(99,102,241,0.3);
        transform: translateY(-2px);
    }
    .cap-icon  { font-size: 1.6rem; margin-bottom: 0.6rem; }
    .cap-title { font-size: 0.95rem; font-weight: 700; color: #e2e8f0; margin-bottom: 0.2rem; }
    .cap-agents{ font-size: 0.75rem; color: #475569; margin-bottom: 0.75rem; font-family: 'JetBrains Mono', monospace; }
    .cap-tags  { display: flex; flex-wrap: wrap; gap: 0.3rem; }
    .cap-tag   {
        background: rgba(99,102,241,0.08);
        border: 1px solid rgba(99,102,241,0.18);
        color: #818cf8;
        font-size: 0.68rem;
        font-weight: 600;
        padding: 0.18rem 0.55rem;
        border-radius: 999px;
    }

    /* ── How it works ─────────────────────────────────────────────────── */
    .how-it-works {
        background: #0d1424;
        border: 1px solid rgba(255,255,255,0.05);
        border-radius: 14px;
        padding: 1.5rem 2rem;
        margin-bottom: 0.5rem;
    }
    .hiw-label {
        font-size: 0.62rem;
        font-weight: 700;
        letter-spacing: 0.14em;
        color: #334155;
        margin-bottom: 1.1rem;
    }
    .hiw-steps {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        flex-wrap: wrap;
    }
    .hiw-step { display: flex; align-items: flex-start; gap: 0.65rem; flex: 1; min-width: 140px; }
    .hiw-num  {
        font-size: 1.1rem;
        font-weight: 800;
        color: rgba(99,102,241,0.35);
        font-family: 'JetBrains Mono', monospace;
        line-height: 1;
        padding-top: 0.05rem;
        flex-shrink: 0;
    }
    .hiw-text strong { display: block; font-size: 0.82rem; color: #cbd5e1; font-weight: 600; margin-bottom: 0.2rem; }
    .hiw-text span   { font-size: 0.75rem; color: #475569; line-height: 1.4; }
    .hiw-arrow { color: #1e293b; font-size: 1.1rem; flex-shrink: 0; }

    /* ── Step headers ─────────────────────────────────────────────────── */
    .step-header {
        display: flex;
        align-items: center;
        gap: 1rem;
        margin: 0.25rem 0 0.85rem;
    }
    .step-num {
        font-size: 0.7rem;
        font-weight: 800;
        font-family: 'JetBrains Mono', monospace;
        background: rgba(99,102,241,0.12);
        border: 1px solid rgba(99,102,241,0.25);
        color: #818cf8;
        padding: 0.25rem 0.55rem;
        border-radius: 6px;
        letter-spacing: 0.05em;
        flex-shrink: 0;
    }
    .step-title { font-size: 1.05rem; font-weight: 700; color: #e2e8f0; margin-bottom: 0.1rem; }
    .step-hint  { font-size: 0.78rem; color: #475569; }

    /* ── Inline hints ─────────────────────────────────────────────────── */
    .inline-hint {
        font-size: 0.78rem;
        font-weight: 500;
        padding: 0.45rem 0.85rem;
        border-radius: 8px;
        margin-top: 0.5rem;
        display: inline-block;
    }
    .inline-hint.warn { background: rgba(245,158,11,0.08); border: 1px solid rgba(245,158,11,0.2); color: #f59e0b; }
    .inline-hint.ok   { background: rgba(16,185,129,0.08); border: 1px solid rgba(16,185,129,0.2);  color: #10b981; }

    /* ── Dispatch card badge ──────────────────────────────────────────── */
    .dispatch-badge {
        display: inline-block;
        font-size: 0.62rem;
        font-weight: 700;
        letter-spacing: 0.1em;
        padding: 0.2rem 0.65rem;
        border-radius: 999px;
        margin-bottom: 0.6rem;
    }

    /* ── Word count widget ────────────────────────────────────────────── */
    .word-count-widget {
        background: rgba(99,102,241,0.06);
        border: 1px solid rgba(99,102,241,0.15);
        border-radius: 12px;
        padding: 1rem;
        text-align: center;
        margin-top: 0.5rem;
    }
    .wcw-val   { font-size: 2rem; font-weight: 800; color: #818cf8; font-family: 'JetBrains Mono', monospace; letter-spacing: -1px; }
    .wcw-label { font-size: 0.75rem; font-weight: 600; color: #64748b; margin-top: 0.1rem; }
    .wcw-sub   { font-size: 0.67rem; color: #334155; margin-top: 0.2rem; }

    /* ── Upload feedback panel ────────────────────────────────────────── */
    .upload-empty {
        text-align: center;
        padding: 1.1rem 0.75rem;
        background: rgba(255,255,255,0.02);
        border: 1px dashed rgba(255,255,255,0.06);
        border-radius: 10px;
        margin: 0.25rem 0.5rem 0.75rem;
    }
    .upload-empty-icon { font-size: 1.4rem; opacity: 0.4; margin-bottom: 0.3rem; }
    .upload-empty-text { font-size: 0.78rem; color: #334155; font-weight: 500; }
    .upload-empty-hint { font-size: 0.68rem; color: #1e293b; margin-top: 0.15rem; }

    .upload-feedback {
        background: rgba(255,255,255,0.02);
        border: 1px solid rgba(255,255,255,0.06);
        border-radius: 10px;
        padding: 0.85rem 0.9rem 0.7rem;
        margin: 0.25rem 0.5rem 0.5rem;
    }
    .upload-count-row {
        display: flex;
        align-items: center;
        gap: 0.5rem;
        margin-bottom: 0.65rem;
    }
    .upload-count { font-size: 1.6rem; font-weight: 800; font-family: 'JetBrains Mono', monospace; line-height: 1; }
    .upload-count-label { font-size: 0.75rem; color: #64748b; font-weight: 500; flex: 1; }
    .upload-quality-badge {
        font-size: 0.62rem; font-weight: 700;
        padding: 0.18rem 0.55rem; border-radius: 999px;
        letter-spacing: 0.05em;
    }
    .upload-file-list { display: flex; flex-direction: column; gap: 0.3rem; }
    .upload-file-row {
        display: flex;
        align-items: center;
        gap: 0.4rem;
        background: rgba(255,255,255,0.03);
        border-radius: 6px;
        padding: 0.28rem 0.5rem;
    }
    .upload-file-ext  { font-size: 0.6rem; font-weight: 800; letter-spacing: 0.04em; flex-shrink: 0; }
    .upload-file-name { font-size: 0.73rem; color: #94a3b8; flex: 1; overflow: hidden; white-space: nowrap; text-overflow: ellipsis; }
    .upload-file-size { font-size: 0.67rem; color: #334155; flex-shrink: 0; }
    .upload-tip {
        font-size: 0.72rem; color: #475569;
        padding: 0.3rem 0.5rem 0.25rem;
        font-style: italic;
    }


    ::-webkit-scrollbar { width: 6px; height: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb { background: #1e293b; border-radius: 3px; }
    ::-webkit-scrollbar-thumb:hover { background: #334155; }

    /* ── LANDING PAGE ─────────────────────────────────────────────────── */
    .landing-wrap {
        min-height: 100vh;
        background: #05070d;
        display: flex;
        flex-direction: column;
        align-items: center;
        position: relative;
        overflow: hidden;
        padding-bottom: 4rem;
    }
    /* top scan-line texture */
    .landing-wrap::before {
        content: '';
        position: fixed;
        inset: 0;
        background-image: repeating-linear-gradient(
            0deg,
            transparent,
            transparent 2px,
            rgba(0,0,0,0.08) 2px,
            rgba(0,0,0,0.08) 4px
        );
        pointer-events: none;
        z-index: 0;
    }
    .landing-bg-glow {
        position: absolute;
        border-radius: 50%;
        filter: blur(120px);
        pointer-events: none;
    }
    .landing-bg-glow-1 {
        width: 700px; height: 700px;
        background: radial-gradient(circle, rgba(99,102,241,0.13) 0%, transparent 70%);
        top: -200px; left: 50%;
        transform: translateX(-50%);
    }
    .landing-bg-glow-2 {
        width: 400px; height: 400px;
        background: radial-gradient(circle, rgba(236,72,153,0.07) 0%, transparent 70%);
        top: 200px; right: -100px;
    }
    .landing-bg-glow-3 {
        width: 350px; height: 350px;
        background: radial-gradient(circle, rgba(6,182,212,0.06) 0%, transparent 70%);
        top: 300px; left: -100px;
    }
    /* top tag strip */
    .landing-tag-strip {
        width: 100%;
        display: flex;
        align-items: center;
        justify-content: space-between;
        padding: 1.2rem 2.5rem;
        position: relative;
        z-index: 2;
        border-bottom: 1px solid rgba(255,255,255,0.03);
    }
    .landing-tag-left {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.6rem;
        letter-spacing: 0.25em;
        color: rgba(99,102,241,0.6);
        text-transform: uppercase;
    }
    .landing-tag-right {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.6rem;
        letter-spacing: 0.2em;
        color: rgba(255,255,255,0.12);
        text-transform: uppercase;
    }
    /* mascot SVG container */
    .mascot-stage {
        position: relative;
        z-index: 2;
        margin: 1.5rem auto 0;
        width: 100%;
        max-width: 900px;
        display: flex;
        justify-content: center;
    }
    .mascot-stage svg {
        width: 100%;
        height: auto;
        filter: drop-shadow(0 0 60px rgba(99,102,241,0.25)) drop-shadow(0 0 120px rgba(236,72,153,0.1));
    }
    /* title block below mascot */
    .landing-title-block {
        position: relative;
        z-index: 2;
        text-align: center;
        margin-top: -1rem;
        padding: 0 1rem;
    }
    .landing-title-block h1 {
        font-family: 'JetBrains Mono', monospace;
        font-size: clamp(2.2rem, 5vw, 4rem);
        font-weight: 700;
        color: #f1f5f9;
        letter-spacing: -0.02em;
        line-height: 1;
        margin: 0 0 0.5rem;
    }
    .landing-title-block h1 span.accent {
        color: transparent;
        -webkit-text-stroke: 1px rgba(99,102,241,0.8);
    }
    .landing-title-block .tagline {
        font-size: 0.82rem;
        color: rgba(148,163,184,0.55);
        letter-spacing: 0.18em;
        text-transform: uppercase;
        font-family: 'JetBrains Mono', monospace;
        margin-bottom: 2.5rem;
    }
    /* discipline pills */
    .landing-pills {
        display: flex;
        gap: 0.6rem;
        justify-content: center;
        flex-wrap: wrap;
        margin-bottom: 3rem;
        position: relative;
        z-index: 2;
    }
    .landing-pill {
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 2px;
        padding: 0.4rem 1rem;
        font-size: 0.68rem;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        font-family: 'JetBrains Mono', monospace;
        color: rgba(148,163,184,0.5);
        transition: all 0.2s;
    }
    .landing-pill.active {
        background: rgba(99,102,241,0.12);
        border-color: rgba(99,102,241,0.35);
        color: rgba(165,180,252,0.9);
    }
    /* auth card on landing */
    .landing-auth-card {
        position: relative;
        z-index: 2;
        background: rgba(13,20,36,0.85);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 1px solid rgba(99,102,241,0.2);
        border-radius: 4px;
        padding: 2.5rem 2.5rem 2rem;
        width: 100%;
        max-width: 380px;
        text-align: center;
        box-shadow: 0 40px 80px rgba(0,0,0,0.5), 0 0 0 1px rgba(99,102,241,0.08);
    }
    .landing-auth-card::before {
        content: '';
        position: absolute;
        top: 0; left: 10%; right: 10%;
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(165,180,252,0.5), transparent);
    }
    .landing-auth-logo {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.62rem;
        letter-spacing: 0.3em;
        color: rgba(165,180,252,0.4);
        text-transform: uppercase;
        margin-bottom: 0.5rem;
    }
    .landing-auth-title {
        font-size: 1.5rem;
        font-weight: 700;
        color: #f1f5f9;
        letter-spacing: -0.5px;
        margin-bottom: 0.25rem;
        font-family: 'Inter', sans-serif;
    }
    .landing-auth-sub {
        font-size: 0.72rem;
        color: rgba(71,85,105,0.8);
        letter-spacing: 0.08em;
        text-transform: uppercase;
        font-family: 'JetBrains Mono', monospace;
        margin-bottom: 1.75rem;
    }
    /* bottom grid line decoration */
    .landing-grid-line {
        position: absolute;
        bottom: 0; left: 0; right: 0;
        height: 180px;
        background:
            linear-gradient(to top, rgba(5,7,13,1) 0%, transparent 100%),
            repeating-linear-gradient(90deg, rgba(99,102,241,0.04) 0px, rgba(99,102,241,0.04) 1px, transparent 1px, transparent 60px),
            repeating-linear-gradient(0deg,  rgba(99,102,241,0.04) 0px, rgba(99,102,241,0.04) 1px, transparent 1px, transparent 60px);
        pointer-events: none;
        z-index: 1;
    }

    </style>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  AUTH
# ─────────────────────────────────────────────

def auth_gate():
    if st.session_state.get("authenticated"):
        return True

    st.markdown("""
    <div class="landing-wrap">

        <!-- atmospheric glows -->
        <div class="landing-bg-glow landing-bg-glow-1"></div>
        <div class="landing-bg-glow landing-bg-glow-2"></div>
        <div class="landing-bg-glow landing-bg-glow-3"></div>
        <div class="landing-grid-line"></div>

        <!-- top strip -->
        <div class="landing-tag-strip">
            <span class="landing-tag-left">Agent43 · System v2.0</span>
            <span class="landing-tag-right">Academic · Intelligence · Engine</span>
        </div>

        <!-- ══════════════ MASCOT SVG ══════════════ -->
        <div class="mascot-stage">
          <svg viewBox="0 0 900 560" xmlns="http://www.w3.org/2000/svg" style="overflow:visible">
            <defs>
              <!-- neon glow filters -->
              <filter id="glow-violet" x="-40%" y="-40%" width="180%" height="180%">
                <feGaussianBlur stdDeviation="6" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
              <filter id="glow-cyan" x="-40%" y="-40%" width="180%" height="180%">
                <feGaussianBlur stdDeviation="5" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
              <filter id="glow-pink" x="-40%" y="-40%" width="180%" height="180%">
                <feGaussianBlur stdDeviation="8" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
              <filter id="glow-soft" x="-60%" y="-60%" width="220%" height="220%">
                <feGaussianBlur stdDeviation="18" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
              <filter id="noise">
                <feTurbulence type="fractalNoise" baseFrequency="0.9" numOctaves="4" stitchTiles="stitch"/>
                <feColorMatrix type="saturate" values="0"/>
                <feBlend in="SourceGraphic" mode="overlay" result="blend"/>
                <feComposite in="blend" in2="SourceGraphic" operator="in"/>
              </filter>
              <!-- gradients -->
              <linearGradient id="body-grad" x1="0" y1="0" x2="0" y2="1">
                <stop offset="0%" stop-color="#1e1b4b"/>
                <stop offset="100%" stop-color="#0f0c24"/>
              </linearGradient>
              <linearGradient id="head-grad" x1="0" y1="0" x2="0" y2="1">
                <stop offset="0%" stop-color="#2d2a5e"/>
                <stop offset="100%" stop-color="#1a1740"/>
              </linearGradient>
              <linearGradient id="screen-grad" x1="0" y1="0" x2="0" y2="1">
                <stop offset="0%" stop-color="#0e7490"/>
                <stop offset="100%" stop-color="#0c4a6e"/>
              </linearGradient>
              <radialGradient id="eye-glow-l" cx="50%" cy="50%" r="50%">
                <stop offset="0%" stop-color="#a5f3fc" stop-opacity="1"/>
                <stop offset="60%" stop-color="#06b6d4" stop-opacity="0.8"/>
                <stop offset="100%" stop-color="#0891b2" stop-opacity="0"/>
              </radialGradient>
              <radialGradient id="eye-glow-r" cx="50%" cy="50%" r="50%">
                <stop offset="0%" stop-color="#f0abfc" stop-opacity="1"/>
                <stop offset="60%" stop-color="#e879f9" stop-opacity="0.8"/>
                <stop offset="100%" stop-color="#a21caf" stop-opacity="0"/>
              </radialGradient>
              <linearGradient id="arm-l-grad" x1="0" y1="0" x2="1" y2="0">
                <stop offset="0%" stop-color="#312e81"/>
                <stop offset="100%" stop-color="#1e1b4b"/>
              </linearGradient>
              <linearGradient id="arm-r-grad" x1="0" y1="0" x2="1" y2="0">
                <stop offset="0%" stop-color="#1e1b4b"/>
                <stop offset="100%" stop-color="#312e81"/>
              </linearGradient>
              <linearGradient id="spray-grad" x1="0" y1="0" x2="1" y2="1">
                <stop offset="0%" stop-color="#c026d3"/>
                <stop offset="100%" stop-color="#7c3aed"/>
              </linearGradient>
              <radialGradient id="spray-burst" cx="30%" cy="30%" r="70%">
                <stop offset="0%" stop-color="#f0abfc" stop-opacity="0.9"/>
                <stop offset="50%" stop-color="#c026d3" stop-opacity="0.5"/>
                <stop offset="100%" stop-color="#7c3aed" stop-opacity="0"/>
              </radialGradient>
            </defs>

            <!-- ─── GRAFFITI TAG BEHIND ROBOT ─── -->
            <!-- Big stylised "43" tag on the wall -->
            <g opacity="0.18" filter="url(#glow-pink)">
              <text x="110" y="440" font-family="'Arial Black', sans-serif" font-size="310" font-weight="900"
                    fill="none" stroke="#c026d3" stroke-width="6" stroke-linejoin="round"
                    letter-spacing="-20">43</text>
            </g>
            <!-- faded outline double -->
            <g opacity="0.07">
              <text x="116" y="447" font-family="'Arial Black', sans-serif" font-size="310" font-weight="900"
                    fill="none" stroke="#06b6d4" stroke-width="3" stroke-linejoin="round"
                    letter-spacing="-20">43</text>
            </g>

            <!-- spray splatter dots behind -->
            <g opacity="0.25" filter="url(#glow-violet)">
              <circle cx="170" cy="160" r="4" fill="#a78bfa"/>
              <circle cx="155" cy="178" r="2.5" fill="#c4b5fd"/>
              <circle cx="185" cy="172" r="3" fill="#818cf8"/>
              <circle cx="145" cy="195" r="2" fill="#a78bfa"/>
              <circle cx="200" cy="155" r="2" fill="#c4b5fd"/>
              <circle cx="162" cy="148" r="1.5" fill="#818cf8"/>
            </g>

            <!-- ─── GROUND SHADOW ─── -->
            <ellipse cx="450" cy="530" rx="230" ry="20" fill="rgba(99,102,241,0.12)" filter="url(#glow-soft)"/>

            <!-- ─── LEGS ─── -->
            <!-- Left leg -->
            <rect x="368" y="445" width="58" height="75" rx="6" fill="url(#body-grad)" stroke="#4f46e5" stroke-width="1.5"/>
            <rect x="375" y="448" width="44" height="8" rx="3" fill="rgba(99,102,241,0.3)"/>
            <!-- left knee joint -->
            <rect x="365" y="490" width="64" height="14" rx="7" fill="#1e1b4b" stroke="#6366f1" stroke-width="1"/>
            <!-- left foot -->
            <rect x="356" y="516" width="76" height="22" rx="8" fill="#0f0c24" stroke="#4f46e5" stroke-width="1.5"/>
            <rect x="360" y="519" width="68" height="6" rx="3" fill="rgba(99,102,241,0.2)"/>

            <!-- Right leg -->
            <rect x="474" y="445" width="58" height="75" rx="6" fill="url(#body-grad)" stroke="#4f46e5" stroke-width="1.5"/>
            <rect x="481" y="448" width="44" height="8" rx="3" fill="rgba(99,102,241,0.3)"/>
            <!-- right knee joint -->
            <rect x="471" y="490" width="64" height="14" rx="7" fill="#1e1b4b" stroke="#6366f1" stroke-width="1"/>
            <!-- right foot -->
            <rect x="468" y="516" width="76" height="22" rx="8" fill="#0f0c24" stroke="#4f46e5" stroke-width="1.5"/>
            <rect x="472" y="519" width="68" height="6" rx="3" fill="rgba(99,102,241,0.2)"/>

            <!-- ─── TORSO ─── -->
            <rect x="330" y="255" width="240" height="200" rx="18" fill="url(#body-grad)" stroke="#4f46e5" stroke-width="2"/>
            <!-- torso panel lines -->
            <rect x="340" y="268" width="220" height="5" rx="2.5" fill="rgba(99,102,241,0.25)"/>
            <!-- chest screen -->
            <rect x="355" y="290" width="190" height="110" rx="10" fill="url(#screen-grad)" stroke="#06b6d4" stroke-width="1.5" opacity="0.9"/>
            <!-- screen scanlines -->
            <rect x="355" y="290" width="190" height="110" rx="10" fill="url(#noise)" opacity="0.05"/>
            <!-- screen content: scrolling code lines -->
            <text x="365" y="310" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.85)">analyzing_brief(context)...</text>
            <text x="365" y="323" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.6)">semantic_retrieval → OK</text>
            <text x="365" y="336" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.7)">agent_dispatch: ALPHA</text>
            <text x="365" y="349" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.5)">citation_index: 14 sources</text>
            <text x="365" y="362" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.75)">writing_stage_3 ████████░░</text>
            <text x="365" y="375" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.4)">originality_score: 91%</text>
            <text x="365" y="388" font-family="'JetBrains Mono',monospace" font-size="8" fill="rgba(165,243,252,0.6)">export → .docx ✓</text>
            <!-- screen cursor blink element -->
            <rect x="365" y="395" width="6" height="9" rx="1" fill="#a5f3fc" opacity="0.9"/>
            <!-- screen glare -->
            <rect x="357" y="292" width="90" height="30" rx="6" fill="rgba(255,255,255,0.04)"/>

            <!-- torso bolts -->
            <circle cx="345" cy="285" r="5" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>
            <circle cx="555" cy="285" r="5" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>
            <circle cx="345" cy="440" r="5" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>
            <circle cx="555" cy="440" r="5" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>

            <!-- torso bottom vents -->
            <rect x="360" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <rect x="384" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <rect x="408" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <rect x="432" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <rect x="456" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <rect x="480" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <rect x="504" y="415" width="18" height="24" rx="3" fill="#0a0e1a" stroke="#4f46e5" stroke-width="1" opacity="0.8"/>
            <!-- vent glow -->
            <rect x="360" y="438" width="162" height="4" rx="2" fill="rgba(99,102,241,0.4)" filter="url(#glow-violet)"/>

            <!-- shoulder joints -->
            <circle cx="330" cy="280" r="22" fill="#1e1b4b" stroke="#6366f1" stroke-width="2"/>
            <circle cx="330" cy="280" r="12" fill="#312e81" stroke="#818cf8" stroke-width="1.5"/>
            <circle cx="330" cy="280" r="5"  fill="#a78bfa"/>
            <circle cx="570" cy="280" r="22" fill="#1e1b4b" stroke="#6366f1" stroke-width="2"/>
            <circle cx="570" cy="280" r="12" fill="#312e81" stroke="#818cf8" stroke-width="1.5"/>
            <circle cx="570" cy="280" r="5"  fill="#a78bfa"/>

            <!-- ─── LEFT ARM (raised, holding spray can) ─── -->
            <g transform="rotate(-40, 308, 280)">
              <!-- upper arm -->
              <rect x="240" y="262" width="70" height="36" rx="12" fill="url(#arm-l-grad)" stroke="#4f46e5" stroke-width="1.5"/>
              <!-- elbow joint -->
              <circle cx="252" cy="280" r="14" fill="#1e1b4b" stroke="#6366f1" stroke-width="1.5"/>
              <circle cx="252" cy="280" r="7"  fill="#312e81" stroke="#818cf8" stroke-width="1"/>
              <!-- forearm -->
              <rect x="178" y="265" width="78" height="30" rx="10" fill="url(#arm-l-grad)" stroke="#4f46e5" stroke-width="1.5"/>
              <!-- wrist -->
              <circle cx="190" cy="280" r="12" fill="#1e1b4b" stroke="#6366f1" stroke-width="1.5"/>
              <!-- hand / fist holding can -->
              <rect x="148" y="265" width="46" height="32" rx="8" fill="#312e81" stroke="#6366f1" stroke-width="1.5"/>
              <!-- knuckle lines -->
              <line x1="156" y1="274" x2="156" y2="289" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
              <line x1="164" y1="272" x2="164" y2="290" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
              <line x1="172" y1="272" x2="172" y2="290" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
              <line x1="180" y1="274" x2="180" y2="289" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
            </g>

            <!-- ─── SPRAY CAN (separate, positioned at left hand area) ─── -->
            <!-- spray can body -->
            <rect x="90" y="100" width="34" height="80" rx="10" fill="url(#spray-grad)" stroke="#c026d3" stroke-width="1.5" filter="url(#glow-pink)"/>
            <!-- can label band -->
            <rect x="90" y="125" width="34" height="28" fill="rgba(240,171,252,0.2)"/>
            <text x="107" y="143" font-family="'Arial Black',sans-serif" font-size="9" font-weight="900" fill="white" text-anchor="middle" opacity="0.9">43</text>
            <!-- can cap -->
            <rect x="95" y="92" width="24" height="12" rx="5" fill="#e879f9" stroke="#f0abfc" stroke-width="1"/>
            <!-- can nozzle -->
            <rect x="107" y="84" width="8" height="12" rx="3" fill="#f5d0fe"/>
            <!-- spray burst emanating from nozzle -->
            <g filter="url(#glow-pink)" opacity="0.85">
              <!-- main burst cloud -->
              <ellipse cx="95" cy="72" rx="28" ry="20" fill="url(#spray-burst)" opacity="0.7"/>
              <ellipse cx="78"  cy="58" rx="18" ry="13" fill="rgba(192,38,211,0.5)"/>
              <ellipse cx="115" cy="60" rx="15" ry="10" fill="rgba(124,58,237,0.5)"/>
              <!-- fine mist dots -->
              <circle cx="60"  cy="52" r="3.5" fill="#f0abfc" opacity="0.7"/>
              <circle cx="72"  cy="42" r="2.5" fill="#e879f9" opacity="0.6"/>
              <circle cx="85"  cy="36" r="2"   fill="#c026d3" opacity="0.5"/>
              <circle cx="100" cy="38" r="3"   fill="#a78bfa" opacity="0.6"/>
              <circle cx="115" cy="44" r="2"   fill="#f0abfc" opacity="0.5"/>
              <circle cx="50"  cy="64" r="2"   fill="#e879f9" opacity="0.4"/>
              <circle cx="126" cy="55" r="2.5" fill="#c4b5fd" opacity="0.5"/>
              <circle cx="55"  cy="45" r="1.5" fill="#f5d0fe" opacity="0.4"/>
              <circle cx="130" cy="43" r="1.5" fill="#f0abfc" opacity="0.3"/>
            </g>

            <!-- graffiti text being sprayed -->
            <g filter="url(#glow-pink)" opacity="0.9">
              <!-- "AGENT" tag in wild style -->
              <text x="155" y="65" font-family="'Arial Black',sans-serif" font-size="38" font-weight="900"
                    fill="none" stroke="#c026d3" stroke-width="3.5" stroke-linejoin="round"
                    transform="rotate(-8, 155, 65)" letter-spacing="2">AGENT</text>
              <text x="157" y="65" font-family="'Arial Black',sans-serif" font-size="38" font-weight="900"
                    fill="rgba(192,38,211,0.35)" transform="rotate(-8, 155, 65)" letter-spacing="2">AGENT</text>
            </g>

            <!-- ─── RIGHT ARM (relaxed, slightly extended) ─── -->
            <g transform="rotate(20, 592, 280)">
              <!-- upper arm -->
              <rect x="590" y="262" width="70" height="36" rx="12" fill="url(#arm-r-grad)" stroke="#4f46e5" stroke-width="1.5"/>
              <!-- elbow -->
              <circle cx="648" cy="280" r="14" fill="#1e1b4b" stroke="#6366f1" stroke-width="1.5"/>
              <circle cx="648" cy="280" r="7"  fill="#312e81" stroke="#818cf8" stroke-width="1"/>
              <!-- forearm -->
              <rect x="644" y="265" width="78" height="30" rx="10" fill="url(#arm-r-grad)" stroke="#4f46e5" stroke-width="1.5"/>
              <!-- wrist -->
              <circle cx="710" cy="280" r="12" fill="#1e1b4b" stroke="#6366f1" stroke-width="1.5"/>
              <!-- open hand / fingers pointing outward -->
              <rect x="706" y="263" width="46" height="34" rx="8" fill="#312e81" stroke="#6366f1" stroke-width="1.5"/>
              <!-- finger lines -->
              <line x1="714" y1="271" x2="714" y2="288" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
              <line x1="722" y1="269" x2="722" y2="290" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
              <line x1="730" y1="269" x2="730" y2="290" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
              <line x1="738" y1="271" x2="738" y2="288" stroke="#4f46e5" stroke-width="1.5" stroke-linecap="round"/>
            </g>

            <!-- ─── NECK ─── -->
            <rect x="414" y="218" width="72" height="42" rx="6" fill="#1a1740" stroke="#4f46e5" stroke-width="1.5"/>
            <!-- neck segments -->
            <rect x="414" y="227" width="72" height="5" rx="2" fill="rgba(99,102,241,0.3)"/>
            <rect x="414" y="238" width="72" height="5" rx="2" fill="rgba(99,102,241,0.25)"/>
            <rect x="414" y="249" width="72" height="5" rx="2" fill="rgba(99,102,241,0.2)"/>

            <!-- ─── HEAD ─── -->
            <rect x="340" y="80" width="220" height="145" rx="22" fill="url(#head-grad)" stroke="#6366f1" stroke-width="2.5"/>
            <!-- head top detail -->
            <rect x="350" y="90"  width="200" height="6" rx="3" fill="rgba(99,102,241,0.3)"/>
            <!-- antenna left -->
            <rect x="380" y="58" width="8" height="28" rx="4" fill="#4f46e5" stroke="#818cf8" stroke-width="1"/>
            <circle cx="384" cy="52" r="9" fill="#312e81" stroke="#818cf8" stroke-width="1.5" filter="url(#glow-violet)"/>
            <circle cx="384" cy="52" r="5" fill="#a78bfa" filter="url(#glow-violet)"/>
            <!-- antenna right -->
            <rect x="512" y="58" width="8" height="28" rx="4" fill="#4f46e5" stroke="#818cf8" stroke-width="1"/>
            <circle cx="516" cy="52" r="9" fill="#312e81" stroke="#818cf8" stroke-width="1.5" filter="url(#glow-violet)"/>
            <circle cx="516" cy="52" r="5" fill="#a78bfa" filter="url(#glow-violet)"/>
            <!-- head side bolts -->
            <circle cx="345" cy="120" r="6" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>
            <circle cx="555" cy="120" r="6" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>
            <circle cx="345" cy="185" r="6" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>
            <circle cx="555" cy="185" r="6" fill="#1e1b4b" stroke="#4f46e5" stroke-width="1.5"/>

            <!-- ─── FACE PANEL ─── -->
            <rect x="358" y="100" width="184" height="110" rx="14" fill="#0a0e1c" stroke="#4f46e5" stroke-width="1.5"/>

            <!-- EYE LEFT — cyan -->
            <rect x="375" y="118" width="64" height="42" rx="10" fill="#0e1a2e" stroke="#06b6d4" stroke-width="2"/>
            <rect x="378" y="121" width="58" height="36" rx="8" fill="url(#eye-glow-l)" opacity="0.15"/>
            <!-- eye iris -->
            <circle cx="407" cy="139" r="14" fill="#0891b2" filter="url(#glow-cyan)"/>
            <circle cx="407" cy="139" r="9"  fill="#06b6d4"/>
            <circle cx="407" cy="139" r="5"  fill="#a5f3fc"/>
            <!-- eye shine -->
            <circle cx="412" cy="134" r="2.5" fill="white" opacity="0.9"/>
            <!-- eye scan line -->
            <rect x="376" y="137" width="62" height="3" rx="1.5" fill="rgba(6,182,212,0.4)"/>

            <!-- EYE RIGHT — pink/magenta -->
            <rect x="461" y="118" width="64" height="42" rx="10" fill="#1a0e2e" stroke="#e879f9" stroke-width="2"/>
            <rect x="464" y="121" width="58" height="36" rx="8" fill="url(#eye-glow-r)" opacity="0.15"/>
            <!-- eye iris -->
            <circle cx="493" cy="139" r="14" fill="#a21caf" filter="url(#glow-pink)"/>
            <circle cx="493" cy="139" r="9"  fill="#e879f9"/>
            <circle cx="493" cy="139" r="5"  fill="#f0abfc"/>
            <!-- eye shine -->
            <circle cx="498" cy="134" r="2.5" fill="white" opacity="0.9"/>
            <!-- eye scan line -->
            <rect x="462" y="137" width="62" height="3" rx="1.5" fill="rgba(232,121,249,0.4)"/>

            <!-- MOUTH / speaker grille -->
            <rect x="375" y="176" width="150" height="26" rx="8" fill="#07091a" stroke="#4f46e5" stroke-width="1.5"/>
            <!-- speaker slots -->
            <rect x="382" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="397" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="412" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="427" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="442" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="457" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="472" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="487" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <rect x="502" y="180" width="10" height="18" rx="3" fill="#1e1b4b"/>
            <!-- mouth glow line -->
            <rect x="375" y="200" width="150" height="3" rx="1.5" fill="rgba(99,102,241,0.5)" filter="url(#glow-violet)"/>

            <!-- ─── FLOATING TAGS around the robot ─── -->
            <!-- top right tag -->
            <g opacity="0.55" transform="rotate(12, 680, 130)" filter="url(#glow-violet)">
              <text x="620" y="148" font-family="'Arial Black',sans-serif" font-size="22" font-weight="900"
                    fill="none" stroke="#818cf8" stroke-width="2.5" stroke-linejoin="round">WRITE</text>
              <text x="622" y="148" font-family="'Arial Black',sans-serif" font-size="22" font-weight="900"
                    fill="rgba(99,102,241,0.3)">WRITE</text>
            </g>
            <!-- bottom left tag -->
            <g opacity="0.45" transform="rotate(-6, 130, 390)" filter="url(#glow-cyan)">
              <text x="60" y="408" font-family="'Arial Black',sans-serif" font-size="20" font-weight="900"
                    fill="none" stroke="#06b6d4" stroke-width="2" stroke-linejoin="round">CITE</text>
              <text x="62" y="408" font-family="'Arial Black',sans-serif" font-size="20" font-weight="900"
                    fill="rgba(6,182,212,0.25)">CITE</text>
            </g>
            <!-- far right tag -->
            <g opacity="0.4" transform="rotate(-14, 780, 350)" filter="url(#glow-pink)">
              <text x="710" y="368" font-family="'Arial Black',sans-serif" font-size="18" font-weight="900"
                    fill="none" stroke="#e879f9" stroke-width="2" stroke-linejoin="round">THINK</text>
              <text x="712" y="368" font-family="'Arial Black',sans-serif" font-size="18" font-weight="900"
                    fill="rgba(232,121,249,0.25)">THINK</text>
            </g>

            <!-- ─── AMBIENT PARTICLES ─── -->
            <g filter="url(#glow-cyan)" opacity="0.6">
              <circle cx="760" cy="180" r="3" fill="#06b6d4"/>
              <circle cx="780" cy="210" r="2" fill="#a5f3fc"/>
              <circle cx="750" cy="240" r="2.5" fill="#06b6d4"/>
              <circle cx="800" cy="160" r="1.5" fill="#67e8f9"/>
            </g>
            <g filter="url(#glow-violet)" opacity="0.5">
              <circle cx="140" cy="310" r="3"   fill="#818cf8"/>
              <circle cx="118" cy="280" r="2"   fill="#c4b5fd"/>
              <circle cx="155" cy="340" r="2.5" fill="#a78bfa"/>
            </g>
            <g filter="url(#glow-pink)" opacity="0.4">
              <circle cx="820" cy="420" r="3" fill="#f0abfc"/>
              <circle cx="840" cy="390" r="2" fill="#e879f9"/>
              <circle cx="800" cy="445" r="2" fill="#f5d0fe"/>
            </g>

          </svg>
        </div>
        <!-- ════════════════════════════════ -->

        <!-- title block -->
        <div class="landing-title-block">
            <h1>AGENT<span class="accent">43</span></h1>
            <div class="tagline">AI · Academic · Writing · System</div>
        </div>

        <!-- discipline pills -->
        <div class="landing-pills">
            <span class="landing-pill active">International Business</span>
            <span class="landing-pill active">International Marketing</span>
            <span class="landing-pill active">Health &amp; Social Care</span>
            <span class="landing-pill">9 Specialist Agents</span>
            <span class="landing-pill">Zero Hallucinations</span>
        </div>

    </div>
    """, unsafe_allow_html=True)

    # ── Auth card (rendered via Streamlit widgets, below the landing art) ──
    st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1.4, 1, 1.4])
    with col2:
        st.markdown("""
        <div class="landing-auth-card">
            <div class="landing-auth-logo">A · G · E · N · T · 43</div>
            <div class="landing-auth-title">Restricted Access</div>
            <div class="landing-auth-sub">Enter your passkey to continue</div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
        pwd = st.text_input("Password", type="password", label_visibility="collapsed",
                            placeholder="Enter access password...")
        if st.button("Enter System", use_container_width=True):
            if pwd == APP_PASSWORD:
                st.session_state["authenticated"] = True
                st.rerun()
            else:
                st.error("Incorrect password.")
    return False

# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────

def render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-brand">
            <span class="sidebar-logo">⬡</span>
            <span class="sidebar-title">Agent43</span>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

        # Navigation
        page = st.radio("nav", ["✍️  Write", "📊  Dashboard", "📚  History"],
                        label_visibility="collapsed")
        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

        # ── Reference Materials Upload ─────────────────────────────────
        st.markdown('<div class="sidebar-section-label">Reference Materials</div>', unsafe_allow_html=True)

        uploaded_files = st.file_uploader(
            "upload_refs",
            accept_multiple_files=True,
            type=["pdf", "docx", "txt"],
            label_visibility="collapsed",
            help="Upload PDFs, Word docs, or text files. Citations will be drawn only from these.",
            key="sidebar_uploads"
        )

        # Store in session state so page_write can access them
        st.session_state["uploaded_files"] = uploaded_files or []

        # Feedback panel
        n = len(uploaded_files) if uploaded_files else 0
        if n == 0:
            st.markdown("""
            <div class="upload-empty">
                <div class="upload-empty-icon">📂</div>
                <div class="upload-empty-text">No papers uploaded yet</div>
                <div class="upload-empty-hint">PDF · DOCX · TXT</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            badge_color = "#10b981" if n >= 3 else "#f59e0b" if n >= 1 else "#ef4444"
            badge_bg    = "#022c22" if n >= 3 else "#2d1a00" if n >= 1 else "#2d0000"
            quality_label = "Excellent" if n >= 4 else "Good" if n >= 2 else "Minimal"
            st.markdown(f"""
            <div class="upload-feedback">
                <div class="upload-count-row">
                    <span class="upload-count" style="color:{badge_color};">{n}</span>
                    <span class="upload-count-label">paper{"s" if n != 1 else ""} loaded</span>
                    <span class="upload-quality-badge" style="background:{badge_bg};color:{badge_color};">{quality_label}</span>
                </div>
                <div class="upload-file-list">
            """, unsafe_allow_html=True)

            for f in uploaded_files:
                ext  = f.name.rsplit(".", 1)[-1].upper() if "." in f.name else "FILE"
                size = f"{f.size / 1024:.0f} KB" if hasattr(f, "size") else ""
                ext_color = {"PDF": "#f87171", "DOCX": "#60a5fa", "TXT": "#a3e635"}.get(ext, "#94a3b8")
                short_name = f.name[:24] + "…" if len(f.name) > 26 else f.name
                st.markdown(f"""
                    <div class="upload-file-row">
                        <span class="upload-file-ext" style="color:{ext_color};">{ext}</span>
                        <span class="upload-file-name">{short_name}</span>
                        <span class="upload-file-size">{size}</span>
                    </div>
                """, unsafe_allow_html=True)

            st.markdown("""
                </div>
            </div>
            """, unsafe_allow_html=True)

            if n < 2:
                st.markdown('<div class="upload-tip">💡 Add more sources for richer citation synthesis</div>', unsafe_allow_html=True)

        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

        # Total cost
        total = get_total_cost()
        st.markdown(f"""
        <div class="sidebar-cost-block">
            <div class="sidebar-cost-label">Cumulative Spend</div>
            <div class="sidebar-cost-value">${total:.4f} <span class="sidebar-cost-unit">USD</span></div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

        # Agent registry
        st.markdown('<div class="sidebar-section-label">Agent Registry</div>', unsafe_allow_html=True)
        for cls, agents in CLASS_AGENTS.items():
            with st.expander(cls, expanded=False):
                for a in agents:
                    info = AGENTS[a]
                    st.markdown(f"""
                    <div class="agent-card">
                        <div class="agent-card-name">{a}</div>
                        <div class="agent-card-sig">{info['signature']}</div>
                    </div>
                    """, unsafe_allow_html=True)

        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
        if st.button("🔒  Sign Out", use_container_width=True):
            st.session_state.clear()
            st.rerun()

    raw = page.strip()
    return raw.split("  ", 1)[1] if "  " in raw else raw.split(" ", 1)[1]

# ─────────────────────────────────────────────
#  PAGE: WRITE
# ─────────────────────────────────────────────

def page_write():
    uploaded_files = st.session_state.get("uploaded_files", [])
    n_papers       = len(uploaded_files)
    dispatch_result = st.session_state.get("dispatch_result")
    selected_agent  = st.session_state.get("selected_agent")
    has_output      = bool(st.session_state.get("last_output"))

    # ══════════════════════════════════════════════════════════════════
    # HERO — only shown before any output exists
    # ══════════════════════════════════════════════════════════════════
    if not has_output:
        st.markdown("""
        <div class="hero">
            <div class="hero-glow hero-glow-left"></div>
            <div class="hero-glow hero-glow-right"></div>
            <div class="hero-inner">
                <div class="hero-eyebrow">AI-POWERED ACADEMIC WRITING SYSTEM</div>
                <h1 class="hero-title">Write with the precision<br>of a first-class scholar.</h1>
                <p class="hero-subtitle">
                    Nine specialist agents. Three disciplines. Citations sourced exclusively
                    from your uploaded references — zero hallucinations, every time.
                </p>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <div class="capability-grid">
            <div class="cap-card">
                <div class="cap-icon">🌐</div>
                <div class="cap-title">International Business</div>
                <div class="cap-agents">Alpha · Beta · Gamma</div>
                <div class="cap-tags">
                    <span class="cap-tag">FDI Theory</span>
                    <span class="cap-tag">Institutional Economics</span>
                    <span class="cap-tag">Global Strategy</span>
                </div>
            </div>
            <div class="cap-card">
                <div class="cap-icon">📣</div>
                <div class="cap-title">International Marketing</div>
                <div class="cap-agents">Delta · Epsilon · Zeta</div>
                <div class="cap-tags">
                    <span class="cap-tag">Brand Equity</span>
                    <span class="cap-tag">Cultural Intelligence</span>
                    <span class="cap-tag">Consumer Behaviour</span>
                </div>
            </div>
            <div class="cap-card">
                <div class="cap-icon">🏥</div>
                <div class="cap-title">Health &amp; Social Care</div>
                <div class="cap-agents">Eta · Theta · Iota</div>
                <div class="cap-tags">
                    <span class="cap-tag">Reflective Practice</span>
                    <span class="cap-tag">Policy &amp; Law</span>
                    <span class="cap-tag">Social Justice</span>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <div class="how-it-works">
            <div class="hiw-label">HOW IT WORKS</div>
            <div class="hiw-steps">
                <div class="hiw-step">
                    <div class="hiw-num">01</div>
                    <div class="hiw-text">
                        <strong>Upload references</strong>
                        <span>Drop PDFs or Word docs in the sidebar panel</span>
                    </div>
                </div>
                <div class="hiw-arrow">→</div>
                <div class="hiw-step">
                    <div class="hiw-num">02</div>
                    <div class="hiw-text">
                        <strong>Describe your brief</strong>
                        <span>Paste your assessment question below</span>
                    </div>
                </div>
                <div class="hiw-arrow">→</div>
                <div class="hiw-step">
                    <div class="hiw-num">03</div>
                    <div class="hiw-text">
                        <strong>Agent recommends</strong>
                        <span>Dispatcher selects the best specialist</span>
                    </div>
                </div>
                <div class="hiw-arrow">→</div>
                <div class="hiw-step">
                    <div class="hiw-num">04</div>
                    <div class="hiw-text">
                        <strong>Write &amp; assess</strong>
                        <span>Essay streams live with originality report</span>
                    </div>
                </div>
            </div>
        </div>
        <div style="height:1.5rem"></div>
        """, unsafe_allow_html=True)

    # ══════════════════════════════════════════════════════════════════
    # STEP 1 — Assessment Brief
    # ══════════════════════════════════════════════════════════════════
    if not has_output:
        st.markdown("""
        <div class="step-header">
            <span class="step-num">01</span>
            <div>
                <div class="step-title">Assessment Brief</div>
                <div class="step-hint">Paste your question, topic, or full assessment context</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        context = st.text_area(
            "brief", height=160, label_visibility="collapsed",
            placeholder="e.g.  Critically evaluate the role of institutional voids in shaping entry strategies of MNEs in emerging markets. Discuss with reference to relevant theory and empirical evidence..."
        )

        if n_papers == 0:
            st.markdown('<div class="inline-hint warn">⚠️  No references uploaded — add PDFs in the sidebar before generating</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="inline-hint ok">✓  {n_papers} reference{"s" if n_papers!=1 else ""} ready · citations sourced exclusively from your documents</div>', unsafe_allow_html=True)

    # ══════════════════════════════════════════════════════════════════
    # STEP 2 — Dispatcher
    # ══════════════════════════════════════════════════════════════════
    if not has_output:
        if 'context' not in dir():
            context = ""

        if context.strip() and not dispatch_result:
            st.markdown("<div style='margin:1rem 0 0.5rem'></div>", unsafe_allow_html=True)
            if st.button("⚡  Analyse Brief & Recommend Agent"):
                with st.spinner("Dispatcher reading your brief..."):
                    result = run_dispatcher(context)
                    st.session_state["dispatch_result"] = result
                    st.session_state["selected_agent"]  = result.get("recommended", "Agent Alpha")
                    st.rerun()

        if dispatch_result:
            rec = dispatch_result
            conf = rec.get("confidence", "")
            conf_color = {"High": "#10b981", "Medium": "#f59e0b", "Low": "#94a3b8"}.get(conf, "#94a3b8")
            conf_bg    = {"High": "#022c22", "Medium": "#2d1a00", "Low": "#1e293b"}.get(conf, "#1e293b")
            st.markdown(f"""
            <div class="dispatch-card">
                <div class="dispatch-badge" style="background:{conf_bg};color:{conf_color};">{conf.upper()} CONFIDENCE</div>
                <h3>⚡ {rec.get('recommended','N/A')}</h3>
                <p><strong>Discipline:</strong> {rec.get('class','')}</p>
                <p><strong>Rationale:</strong> {rec.get('reasoning','')}</p>
                <p style="opacity:0.55;font-size:0.78rem;margin-top:0.5rem;">{rec.get('not_alpha','')} · {rec.get('not_beta','')}</p>
            </div>
            """, unsafe_allow_html=True)

            all_agents = list(AGENTS.keys())
            override = st.selectbox("Confirm or override agent",
                                    all_agents,
                                    index=all_agents.index(rec.get("recommended", all_agents[0])))
            st.session_state["selected_agent"] = override

            if st.button("🔄  Re-run Dispatcher"):
                st.session_state.pop("dispatch_result", None)
                st.session_state.pop("selected_agent", None)
                st.rerun()

    # ══════════════════════════════════════════════════════════════════
    # STEP 3 — Writing Brief
    # ══════════════════════════════════════════════════════════════════
    if selected_agent and not has_output:
        st.markdown('<hr style="border-color:rgba(255,255,255,0.05);margin:1.75rem 0 1.25rem">', unsafe_allow_html=True)
        agent_info = AGENTS[selected_agent]

        st.markdown(f"""
        <div class="step-header">
            <span class="step-num">02</span>
            <div>
                <div class="step-title">Brief for <span style="color:#a5b4fc">{selected_agent}</span></div>
                <div class="step-hint">{agent_info['class']} · {agent_info['signature']}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([3, 1])
        with col1:
            structure = st.text_area(
                "Essay Structure / Outline", height=130,
                placeholder="e.g.\n1. Introduction\n2. Literature Review\n3. Critical Analysis\n4. Conclusion\n5. References"
            )
            rubric = st.text_area(
                "Marking Rubric (optional)", height=90,
                placeholder="Paste marking criteria here — the agent will align its writing to each criterion..."
            )
        with col2:
            word_count = st.number_input("Target Words", min_value=200, max_value=8000,
                                         value=1500, step=100)
            st.markdown(f"""
            <div class="word-count-widget">
                <div class="wcw-val">{word_count:,}</div>
                <div class="wcw-label">body words</div>
                <div class="wcw-sub">references excluded</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<div style='margin-top:1.25rem'></div>", unsafe_allow_html=True)
        if st.button(f"✍️  Generate with {selected_agent}", use_container_width=False):
            if not context.strip():
                st.error("Please provide an assessment brief.")
                return
            if not structure.strip():
                st.error("Please provide an essay structure.")
                return
            if not uploaded_files:
                st.warning("⚠️ No reference materials uploaded — the agent will write without citations.")

            source_texts = []
            for f in uploaded_files:
                txt = extract_text_from_file(f)
                if txt:
                    source_texts.append(txt)

            stage_status = st.empty()
            stage_status.info("⚙️ Stage 1 — Extracting citation index from sources...")
            st.markdown('<hr style="border-color:rgba(255,255,255,0.05)">', unsafe_allow_html=True)
            st.markdown('<div class="section-label">Live Output</div>', unsafe_allow_html=True)
            stream_box = st.empty()
            stage_status.info("⚙️ Stage 2 — Building section plan...")

            output, tok_in, tok_out, cost = run_writer(
                selected_agent, context, structure,
                rubric, word_count, source_texts,
                stream_container=stream_box
            )
            stage_status.success("✅ Write-up complete.")

            st.session_state["last_output"]       = output
            st.session_state["last_agent"]        = selected_agent
            st.session_state["last_discipline"]   = agent_info["class"]
            st.session_state["last_context"]      = context
            st.session_state["last_structure"]    = structure
            st.session_state["last_word_count"]   = word_count
            st.session_state["last_tokens_in"]    = tok_in
            st.session_state["last_tokens_out"]   = tok_out
            st.session_state["last_cost"]         = cost
            st.session_state["last_source_texts"] = source_texts
            st.session_state["assessment_done"]   = False
            st.rerun()

    # ── OUTPUT + ASSESSMENT ────────────────────────────────────────────
    if st.session_state.get("last_output"):
        output       = st.session_state["last_output"]
        agent_name   = st.session_state["last_agent"]
        discipline   = st.session_state["last_discipline"]
        source_texts = st.session_state.get("last_source_texts", [])
        structure    = st.session_state.get("last_structure", "")

        # Clean output — strip markdown symbols
        cleaned_output = clean_output_text(output)
        body_words     = count_body_words(cleaned_output)
        total_words    = len(cleaned_output.split())

        st.markdown("---")
        st.markdown(
            f'<div class="section-label">Generated Write-Up</div>',
            unsafe_allow_html=True
        )
        st.markdown(
            f'<div class="output-meta">{agent_name} &nbsp;·&nbsp; {discipline} &nbsp;·&nbsp; '
            f'Body: {body_words:,} words &nbsp;·&nbsp; Total (inc. references): {total_words:,} words</div>',
            unsafe_allow_html=True
        )

        with st.expander("📄 Read / Review Full Write-Up", expanded=False):
            # Render paragraphs cleanly — no markdown bleed-through
            paragraphs = [p.strip() for p in cleaned_output.split("\n\n") if p.strip()]
            rendered = ""
            for para in paragraphs:
                # Treat short lines (likely headings) differently
                if len(para) < 100 and not para.endswith("."):
                    rendered += f'<p class="write-heading">{para}</p>'
                else:
                    rendered += f'<p class="write-para">{para}</p>'
            st.markdown(f'<div class="writing-output">{rendered}</div>', unsafe_allow_html=True)

        # Assessment
        st.markdown("---")
        st.markdown('<div class="section-label">Assessment Report</div>', unsafe_allow_html=True)

        if not st.session_state.get("assessment_done"):
            if st.button("🔍 Run Assessment", use_container_width=False):
                with st.spinner("Running similarity checks, risk assessment, and originality analysis..."):
                    src_sim      = check_source_similarity(cleaned_output, source_texts) if source_texts else None
                    hist_sim     = check_history_similarity(cleaned_output)
                    risk         = run_risk_assessment(cleaned_output)
                    originality  = run_originality_score(cleaned_output, source_texts)
                    cite_check   = run_citation_verification(cleaned_output)
                    save_writing(
                        agent_name, discipline,
                        st.session_state["last_context"],
                        body_words, cleaned_output,
                        st.session_state["last_tokens_in"],
                        st.session_state["last_tokens_out"],
                        st.session_state["last_cost"]
                    )
                st.session_state["assess_src_sim"]    = src_sim
                st.session_state["assess_hist_sim"]   = hist_sim
                st.session_state["assess_risk"]       = risk
                st.session_state["assess_originality"]= originality
                st.session_state["assess_cite_check"] = cite_check
                st.session_state["assess_cost"]       = st.session_state["last_cost"]
                st.session_state["assessment_done"]   = True
                st.rerun()

        if st.session_state.get("assessment_done"):
            src_sim     = st.session_state["assess_src_sim"]
            hist_sim    = st.session_state["assess_hist_sim"]
            risk        = st.session_state["assess_risk"]
            originality = st.session_state.get("assess_originality", {"score": 0, "label": "N/A", "color": "grey"})
            cite_check  = st.session_state.get("assess_cite_check", {})
            cost        = st.session_state["assess_cost"]
            total       = get_total_cost()

            risk_level = risk.get("risk_level", "Unknown")
            risk_score = risk.get("score", 0)
            risk_color = {"Low": "#10b981", "Medium": "#f59e0b", "High": "#ef4444"}.get(risk_level, "#94a3b8")
            risk_bg    = {"Low": "#022c22", "Medium": "#2d1a00", "High": "#2d0000"}.get(risk_level, "#1e293b")

            src_pct   = f"{src_sim*100:.1f}%" if src_sim is not None else "N/A"
            src_label = sim_band(src_sim)[0] if src_sim is not None else "No sources uploaded"
            src_color = sim_band(src_sim)[1] if src_sim is not None else "grey"
            src_hex   = {"green": "#10b981", "orange": "#f59e0b", "red": "#ef4444", "grey": "#94a3b8"}.get(src_color, "#94a3b8")
            src_bg    = {"green": "#022c22", "orange": "#2d1a00", "red": "#2d0000", "grey": "#1e293b"}.get(src_color, "#1e293b")

            hist_pct   = f"{hist_sim*100:.1f}%"
            hist_label = sim_band(hist_sim)[0]
            hist_color = sim_band(hist_sim)[1]
            hist_hex   = {"green": "#10b981", "orange": "#f59e0b", "red": "#ef4444"}.get(hist_color, "#94a3b8")
            hist_bg    = {"green": "#022c22", "orange": "#2d1a00", "red": "#2d0000"}.get(hist_color, "#1e293b")

            flags_html = ""
            for f in risk.get("flags", []):
                flags_html += f'<span class="flag-pill">{f}</span> '

            orig_score = originality.get("score", 0)
            orig_label = originality.get("label", "N/A")
            orig_color_name = originality.get("color", "grey")
            orig_hex   = {"green": "#10b981", "orange": "#f59e0b", "red": "#ef4444", "grey": "#94a3b8"}.get(orig_color_name, "#94a3b8")
            orig_bg    = {"green": "#022c22", "orange": "#2d1a00", "red": "#2d0000", "grey": "#1e293b"}.get(orig_color_name, "#1e293b")

            st.markdown(f"""
            <div class="assess-grid">

              <div class="assess-card">
                <div class="assess-card-title">Similarity vs Source Materials</div>
                <div class="assess-score" style="color:{src_hex};">{src_pct}</div>
                <div class="assess-label" style="background:{src_bg}; color:{src_hex};">{src_label}</div>
                <div class="assess-desc">Measures semantic overlap between the write-up and uploaded references. Lower = more genuine synthesis rather than paraphrase.</div>
              </div>

              <div class="assess-card">
                <div class="assess-card-title">Similarity vs Past Work</div>
                <div class="assess-score" style="color:{hist_hex};">{hist_pct}</div>
                <div class="assess-label" style="background:{hist_bg}; color:{hist_hex};">{hist_label}</div>
                <div class="assess-desc">Compares against all previous submissions in your history. Flags recycled arguments across different assignments.</div>
              </div>

              <div class="assess-card">
                <div class="assess-card-title">AI Detection Risk</div>
                <div class="assess-score" style="color:{risk_color};">{risk_score}/100</div>
                <div class="assess-label" style="background:{risk_bg}; color:{risk_color};">{risk_level} Risk</div>
                <div class="assess-desc">{risk.get("summary", "")}</div>
                <div style="margin-top:0.75rem;">{flags_html}</div>
              </div>

              <div class="assess-card">
                <div class="assess-card-title">Originality Score</div>
                <div class="assess-score" style="color:{orig_hex};">{orig_score}/100</div>
                <div class="assess-label" style="background:{orig_bg}; color:{orig_hex};">{orig_label}</div>
                <div class="assess-desc">Measures how far the write-up goes beyond its sources. Rewards synthesis across multiple references and novel argumentation.</div>
              </div>

            </div>

            <div class="cost-summary-bar">
              <span class="cost-summary-label">Generation cost</span>
              <span class="cost-summary-value">${cost:.4f}</span>
              <span class="cost-summary-sep">·</span>
              <span class="cost-summary-label">Cumulative total</span>
              <span class="cost-summary-value">${total:.4f} USD</span>
            </div>
            """, unsafe_allow_html=True)

            # ── Citation Verification Panel ──────────────────────────────
            parity      = cite_check.get("parity", False)
            cite_status = cite_check.get("status", "")
            cite_summary= cite_check.get("summary", "")
            orphaned    = cite_check.get("orphaned_intext", [])
            unused      = cite_check.get("unused_references", [])

            if cite_status not in ("no_references", "parse_error", ""):
                parity_color = "#10b981" if parity else "#ef4444"
                parity_bg    = "#022c22" if parity else "#2d0000"
                parity_label = "✅ Full Parity" if parity else "⚠️ Mismatches Found"

                orphaned_html = ""
                if orphaned:
                    orphaned_html = "<div style='margin-top:0.6rem;'><span style='font-size:0.72rem;font-weight:700;color:#f59e0b;text-transform:uppercase;letter-spacing:0.08em;'>Orphaned In-Text Citations</span>"
                    for item in orphaned:
                        orphaned_html += f"<div class='flag-pill' style='color:#ef4444;border-color:rgba(239,68,68,0.3);margin-top:0.2rem;'>{item.get('citation','')} — {item.get('note','')}</div>"
                    orphaned_html += "</div>"

                unused_html = ""
                if unused:
                    unused_html = "<div style='margin-top:0.6rem;'><span style='font-size:0.72rem;font-weight:700;color:#f59e0b;text-transform:uppercase;letter-spacing:0.08em;'>Unused References</span>"
                    for item in unused:
                        ref_short = item.get('reference', '')[:80] + ('…' if len(item.get('reference','')) > 80 else '')
                        unused_html += f"<div class='flag-pill' style='color:#f59e0b;border-color:rgba(245,158,11,0.3);margin-top:0.2rem;'>{ref_short}</div>"
                    unused_html += "</div>"

                st.markdown(f"""
                <div class="assess-card" style="margin-top:1rem;border-color:rgba({('52,211,153' if parity else '239,68,68')},0.3);">
                  <div class="assess-card-title">Citation Parity Check</div>
                  <div class="assess-label" style="background:{parity_bg}; color:{parity_color}; font-size:0.85rem; padding:0.3rem 1rem;">{parity_label}</div>
                  <div class="assess-desc" style="margin-top:0.6rem;">{cite_summary}</div>
                  {orphaned_html}
                  {unused_html}
                </div>
                """, unsafe_allow_html=True)

            # Export
            st.markdown("---")
            st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)
            col1, col2 = st.columns(2)
            with col1:
                docx_bytes = export_docx(cleaned_output, agent_name, discipline, structure)
                st.download_button(
                    "📄 Download DOCX",
                    data=docx_bytes,
                    file_name=f"agent43_{agent_name.replace(' ','_').lower()}_{datetime.date.today()}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with col2:
                st.download_button(
                    "📋 Download TXT",
                    data=cleaned_output.encode("utf-8"),
                    file_name=f"agent43_{agent_name.replace(' ','_').lower()}_{datetime.date.today()}.txt",
                    mime="text/plain",
                    use_container_width=True
                )

        # New write-up
        st.markdown("---")
        if st.button("🆕 Start New Write-Up"):
            for key in ["last_output","last_agent","last_discipline","last_context",
                        "last_structure","last_word_count","last_tokens_in","last_tokens_out",
                        "last_cost","last_source_texts","assessment_done","assess_src_sim",
                        "assess_hist_sim","assess_risk","assess_cite_check","assess_cost",
                        "dispatch_result","selected_agent"]:
                st.session_state.pop(key, None)
            st.rerun()

# ─────────────────────────────────────────────
#  PAGE: DASHBOARD
# ─────────────────────────────────────────────

def page_dashboard():
    st.markdown("""
    <div class="page-header">
        <div class="page-header-content">
            <h1>📊 Dashboard</h1>
            <p>Usage analytics, cost tracking, and agent performance across all sessions</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    try:
        writings = supabase.table("writings").select("*").order("created_at", desc=True).execute().data
        costs    = supabase.table("cost_log").select("*").execute().data
    except Exception as e:
        st.error(f"Could not load data: {e}")
        return

    total_cost   = sum(float(r["cost_usd"]) for r in costs)
    total_writes = len(writings)
    total_words  = sum(int(r.get("word_count", 0)) for r in writings)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Write-Ups",    total_writes)
    c2.metric("Total Words",        f"{total_words:,}")
    c3.metric("Total Cost (USD)",   f"${total_cost:.4f}")
    c4.metric("Avg Cost / Write",   f"${(total_cost/total_writes if total_writes else 0):.4f}")

    st.markdown("---")

    if writings:
        st.markdown("Agent Usage")
        agent_counts = {}
        for w in writings:
            a = w.get("agent_name", "Unknown")
            agent_counts[a] = agent_counts.get(a, 0) + 1
        for agent, count in sorted(agent_counts.items(), key=lambda x: -x[1]):
            st.markdown(f'<div class="assess-desc" style="margin-bottom:0.4rem;">{agent} — {count} write-up{"s" if count != 1 else ""}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("Cost by Feature")
        feature_costs = {}
        for r in costs:
            f = r.get("feature", "unknown")
            feature_costs[f] = feature_costs.get(f, 0) + float(r["cost_usd"])
        for feat, amt in sorted(feature_costs.items(), key=lambda x: -x[1]):
            st.markdown(f'<div class="assess-desc" style="margin-bottom:0.4rem;">{feat} — ${amt:.5f}</div>', unsafe_allow_html=True)
    else:
        st.info("No write-ups yet. Generate your first one in the Write tab.")

# ─────────────────────────────────────────────
#  PAGE: HISTORY
# ─────────────────────────────────────────────

def page_history():
    st.markdown("""
    <div class="page-header">
        <div class="page-header-content">
            <h1>📚 History</h1>
            <p>All past write-ups with full text preview and re-download</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    try:
        writings = supabase.table("writings").select("*").order("created_at", desc=True).limit(50).execute().data
    except Exception as e:
        st.error(f"Could not load history: {e}")
        return

    if not writings:
        st.info("No write-ups saved yet.")
        return

    for w in writings:
        created = w.get("created_at", "")[:10]
        agent   = w.get("agent_name", "Unknown")
        disc    = w.get("discipline", "")
        wc      = w.get("word_count", 0)
        cost    = float(w.get("cost_usd", 0))
        raw_preview = (w.get("output_text", "")[:300] + "...") if w.get("output_text") else ""
        preview = clean_output_text(raw_preview)
        ctx     = w.get("context", "")[:120]

        with st.expander(f"{created}  ·  {agent}  ·  {disc}  ·  {wc:,} words  ·  ${cost:.5f}"):
            st.caption(f"Brief: {ctx}")
            st.markdown(f'<div class="write-para">{preview}</div>', unsafe_allow_html=True)
            docx_bytes = export_docx(w.get("output_text", ""), agent, disc, "")
            st.download_button(
                "📄 Re-download DOCX",
                data=docx_bytes,
                file_name=f"agent43_{agent.replace(' ','_').lower()}_{created}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{w['id']}"
            )

# ─────────────────────────────────────────────
#  PAGE: SETUP
# ─────────────────────────────────────────────

def page_setup():
    st.markdown("## ⚙️ Setup Guide")

    st.markdown("### Streamlit Secrets")
    st.code("""
# .streamlit/secrets.toml or Streamlit Cloud dashboard

OPENAI_API_KEY   = "sk-..."
SUPABASE_URL     = "https://xxxx.supabase.co"
SUPABASE_ANON_KEY = "eyJ..."
APP_PASSWORD     = "your_secure_password"
""", language="toml")

    st.markdown("### Supabase SQL Setup")
    st.markdown("Run this SQL in your **Supabase → SQL Editor**:")
    st.code(SUPABASE_SQL, language="sql")

    st.markdown("### requirements.txt")
    st.code("""streamlit
openai
supabase
tiktoken
numpy
PyPDF2
python-docx
scikit-learn""", language="text")

    st.markdown("### Agent Registry")
    for name, info in AGENTS.items():
        st.markdown(f"**{name}** · {info['class']} · *{info['signature']}*")

# ─────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────

def main():
    inject_css()

    if not auth_gate():
        return

    page = render_sidebar()

    if page == "Write":
        page_write()
    elif page == "Dashboard":
        page_dashboard()
    elif page == "History":
        page_history()

if __name__ == "__main__":
    main()
